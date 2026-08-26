from __future__ import annotations

import hashlib
import io
from datetime import datetime
from pathlib import Path

from docx import Document
from sqlalchemy.orm import Session

from .cip_models import PRODUCT_CIP, PRODUCT_MEP
from .database import SessionLocal
from .models import User
from .services.audit import record
from .sow_models import SOWTemplateVersion

SOW_TEMPLATE_MEP_SMALL_PROJECT = "MEP_SMALL_PROJECT"
SOW_TEMPLATE_CIP_SMALL_PROJECT = "CIP_SMALL_PROJECT"
SMALL_PROJECT_TEMPLATE_KEYS = (
    SOW_TEMPLATE_MEP_SMALL_PROJECT,
    SOW_TEMPLATE_CIP_SMALL_PROJECT,
)

ASSET_DIR = Path(__file__).resolve().parent / "small_project_sow_template_assets"
SMALL_PROJECT_REQUIRED_PLACEHOLDERS = (
    "<CustomerName>",
    "<99999999>",
    "<Today>",
)

_TEMPLATE_META = {
    SOW_TEMPLATE_MEP_SMALL_PROJECT: {
        "label": "MEP Small Project SOW",
        "product_type": PRODUCT_MEP,
        "customer_type": "Install_Base",
        "filename": "MEP_Template_SmallProject_2026_08.docx",
        "sha256": "a075ac54adbdfd1301835a546abbc5a09677b90d6b93e3a084c39b59d8af2226",
    },
    SOW_TEMPLATE_CIP_SMALL_PROJECT: {
        "label": "CIP Small Project SOW",
        "product_type": PRODUCT_CIP,
        "customer_type": "Install_Base",
        "filename": "CIP_Template_SmallProject_2026_07.docx",
        "sha256": "9cd71d907fdf21e4ab51d5f2fda53cd1dafcab6a6c855c87e0ebe0ded8ecbb2a",
    },
}


def small_project_template_meta(template_key: str) -> dict[str, str]:
    try:
        return dict(_TEMPLATE_META[template_key])
    except KeyError as exc:
        raise ValueError(f"Unknown Small Project SOW template key: {template_key}") from exc


def _asset_path(template_key: str) -> Path:
    return ASSET_DIR / small_project_template_meta(template_key)["filename"]


def load_small_project_template_asset(template_key: str) -> bytes:
    """Read and verify the exact controlled Small Project source DOCX.

    The source documents are repository binary assets, not generated or reconstructed
    at runtime. The expected SHA-256 is pinned in source so a missing, replaced, or
    corrupted document fails closed before it can be seeded into the database.
    """
    meta = small_project_template_meta(template_key)
    path = _asset_path(template_key)
    if not path.is_file():
        raise RuntimeError(f"Controlled Small Project SOW template asset is missing: {path.name}")
    content = path.read_bytes()
    digest = hashlib.sha256(content).hexdigest()
    if digest != meta["sha256"]:
        raise RuntimeError(
            f"Controlled Small Project SOW template SHA-256 mismatch for {path.name}: "
            f"expected {meta['sha256']}, got {digest}"
        )
    return content


def _docx_visible_text(content: bytes) -> str:
    try:
        doc = Document(io.BytesIO(content))
    except Exception as exc:
        raise ValueError("The uploaded file is not a valid Word .docx document.") from exc

    chunks: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    for section in doc.sections:
        chunks.extend(p.text for p in section.header.paragraphs)
        chunks.extend(p.text for p in section.footer.paragraphs)
    return "\n".join(chunks)


def validate_small_project_template(content: bytes) -> list[str]:
    text = _docx_visible_text(content)
    return sorted(marker for marker in SMALL_PROJECT_REQUIRED_PLACEHOLDERS if marker not in text)


def seed_small_project_sow_templates(db: Session) -> None:
    admin = db.query(User).filter(User.username_normalized == "admin").first()
    if not admin:
        return

    for template_key in SMALL_PROJECT_TEMPLATE_KEYS:
        if db.query(SOWTemplateVersion).filter(
            SOWTemplateVersion.template_key == template_key
        ).count():
            continue

        meta = small_project_template_meta(template_key)
        content = load_small_project_template_asset(template_key)
        missing = validate_small_project_template(content)
        if missing:
            raise RuntimeError(
                f"Bundled {meta['label']} template is missing required placeholder(s): "
                + ", ".join(missing)
            )

        row = SOWTemplateVersion(
            template_key=template_key,
            label=meta["label"],
            product_type=meta["product_type"],
            customer_type=meta["customer_type"],
            version_no=1,
            status="ACTIVE",
            filename=meta["filename"],
            content=content,
            content_sha256=meta["sha256"],
            change_reason="Initial controlled Small Project SOW source template.",
            created_by=admin.id,
            activated_by=admin.id,
            activated_at=datetime.utcnow(),
        )
        db.add(row)
        db.flush()
        record(
            db,
            event_type="SOW_TEMPLATE_ACTIVATED",
            user_id=admin.id,
            field_name=f"SOW_TEMPLATE:{template_key}:1",
            new_value=row.filename,
            reason=row.change_reason,
        )
    db.commit()


def register_small_project_sow_templates(app) -> None:
    @app.on_event("startup")
    def seed_small_project_sow_templates_on_startup():
        db = SessionLocal()
        try:
            seed_small_project_sow_templates(db)
        finally:
            db.close()
