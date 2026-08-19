from __future__ import annotations

import html
import io
import zipfile

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen.canvas import Canvas
from reportlab.platypus import BaseDocTemplate, Frame, PageBreak, PageTemplate, Paragraph as RP, Spacer, Table, TableStyle
from reportlab.platypus.tableofcontents import TableOfContents

from . import sow_pdf_v3, sow_service
from .cip_sow import docx as cip_docx
from .cip_sow import pdf as cip_pdf
from .sow_models import SOWTemplateVersion
from .sow_signature_layout import compact_signature_spacing, is_signature_table, signature_pdf_flowables


COVER_FOOTER = (
    "This statement of work estimate is the property and proprietary to Data Systems International, Inc. "
    "dba Cloud inventory® and contains trade secret and confidential information and is solely for Customer’s internal use.  "
    "Without the express written consent of Cloud Inventory ®, this estimate shall not be used, reproduced, copied, disclosed, "
    "transmitted in whole or in part.  Copyright © 2026 Data Systems International, Inc. dba Cloud Inventory®.  All rights reserved."
)

_installed = False


def _compact_docx_bytes(raw: bytes) -> bytes:
    doc = Document(io.BytesIO(raw))
    if compact_signature_spacing(doc) == 0:
        return raw
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _logo(docx_bytes: bytes):
    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as archive:
            names = [
                name for name in archive.namelist()
                if name.startswith("word/media/") and name.lower().endswith((".png", ".jpg", ".jpeg"))
            ]
            return archive.read(names[0]) if names else None
    except Exception:
        return None


class _PDFDoc(BaseDocTemplate):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._seq = 0

    def beforeDocument(self):
        self._seq = 0
        return super().beforeDocument()

    def afterFlowable(self, flowable):
        if isinstance(flowable, RP) and hasattr(flowable, "_toc_level"):
            self._seq += 1
            key = f"sow_h_{self._seq}"
            self.canv.bookmarkPage(key)
            self.notify("TOCEntry", (flowable._toc_level, flowable.getPlainText(), self.page, key))


def _canvas(customer: str, estimate: str, logo_bytes: bytes | None):
    class _Canvas(Canvas):
        def __init__(self, *args, **kwargs):
            Canvas.__init__(self, *args, **kwargs)
            self._states = []
            self._logo = ImageReader(io.BytesIO(logo_bytes)) if logo_bytes else None

        def showPage(self):
            self._states.append(dict(self.__dict__))
            self._startPage()

        def save(self):
            total = len(self._states)
            for state in self._states:
                self.__dict__.update(state)
                self._decorate(total)
                Canvas.showPage(self)
            Canvas.save(self)

        def _decorate(self, total):
            width, height = letter
            left = .72 * inch
            right = width - .72 * inch
            if self._logo:
                image_width, image_height = self._logo.getSize()
                draw_width = 120
                draw_height = draw_width * image_height / image_width
                self.drawImage(
                    self._logo,
                    left,
                    height - .28 * inch - draw_height,
                    width=draw_width,
                    height=draw_height,
                    mask="auto",
                    preserveAspectRatio=True,
                )
            if self._pageNumber > 1:
                self.setFont("Helvetica-Bold", 9)
                self.drawRightString(right, height - .82 * inch, customer)
            self.setLineWidth(1)
            self.line(left, height - 1.03 * inch, right, height - 1.03 * inch)
            if self._pageNumber == 1:
                style = ParagraphStyle("coverFooter", fontName="Helvetica", fontSize=6.3, leading=7.2)
                paragraph = RP(html.escape(COVER_FOOTER), style)
                paragraph.wrap(right - left, .6 * inch)
                paragraph.drawOn(self, left, .28 * inch)
            else:
                self.setFont("Helvetica-Bold", 7.5)
                self.drawString(left, .42 * inch, f"Estimate Number: {estimate}")
                self.drawCentredString(width / 2, .42 * inch, f"Page {self._pageNumber} of {total}")

    return _Canvas


def _review_pdf(docx_bytes: bytes, customer: str, estimate: str, sow_date: str, product: str) -> bytes:
    doc = Document(io.BytesIO(docx_bytes))
    out = io.BytesIO()
    frame = Frame(
        .72 * inch,
        .72 * inch,
        letter[0] - 1.44 * inch,
        letter[1] - 1.78 * inch,
        id="body",
        topPadding=0,
        bottomPadding=0,
    )
    pdf = _PDFDoc(out, pagesize=letter)
    pdf.addPageTemplates([PageTemplate(id="sow", frames=[frame])])
    styles = getSampleStyleSheet()

    def add(name, parent, **kwargs):
        styles.add(ParagraphStyle(name, parent=styles[parent], **kwargs))

    add("CoverTitle", "Title", fontName="Helvetica-Bold", fontSize=20, leading=24, alignment=TA_CENTER, spaceAfter=24)
    add("CoverCustomer", "Normal", fontName="Helvetica-Bold", fontSize=16, leading=20, alignment=TA_CENTER, spaceAfter=60)
    add("CoverSubtitle", "Normal", fontName="Helvetica-Bold", fontSize=12.5, leading=16, alignment=TA_CENTER, spaceAfter=22)
    add("CoverMeta", "Normal", fontName="Helvetica-Bold", fontSize=11, leading=14, alignment=TA_CENTER, spaceAfter=28)
    add("TOCTitle", "Heading1", fontName="Helvetica-Bold", fontSize=16, leading=20, alignment=TA_CENTER, spaceAfter=18)
    add("H1S", "Heading1", fontName="Helvetica-Bold", fontSize=14, leading=17, spaceAfter=10, keepWithNext=True)
    add("H2S", "Heading2", fontName="Helvetica-Bold", fontSize=11.5, leading=14, spaceBefore=7, spaceAfter=5, keepWithNext=True)
    add("H3S", "Heading3", fontName="Helvetica-Bold", fontSize=10.5, leading=13, spaceBefore=5, spaceAfter=4, keepWithNext=True)
    add("BodyS", "BodyText", fontName="Helvetica", fontSize=8.7, leading=11.3, spaceAfter=4)
    add("BulletS", "BodyS", leftIndent=18, firstLineIndent=-8, bulletIndent=6, spaceAfter=2)
    add("AppendixS", "Heading1", fontName="Helvetica-Bold", fontSize=16, leading=20, alignment=TA_CENTER, spaceAfter=10, keepWithNext=True)

    subtitle = (
        "Delivery Estimate for Service to Deploy and Configure Cloud Inventory Platform<br/>Solution"
        if product == "CIP"
        else "Delivery Estimate for Service to Deploy and Configure Mobile Enterprise<br/>Platform Solution"
    )
    story = [
        Spacer(1, 1.28 * inch),
        RP("Statement Of Work", styles["CoverTitle"]),
        RP(html.escape(customer), styles["CoverCustomer"]),
        RP(subtitle, styles["CoverSubtitle"]),
        RP(f"Estimate Number: {html.escape(estimate)}", styles["CoverMeta"]),
        RP(f"Date: {html.escape(sow_date)}", styles["CoverMeta"]),
        PageBreak(),
        RP("Table of Contents", styles["TOCTitle"]),
    ]
    toc = TableOfContents()
    toc.levelStyles = [
        ParagraphStyle("T1", fontName="Helvetica-Bold", fontSize=9.2, leading=12, leftIndent=0, spaceBefore=4),
        ParagraphStyle("T2", fontName="Helvetica", fontSize=8.5, leading=11, leftIndent=14),
        ParagraphStyle("T3", fontName="Helvetica", fontSize=8.2, leading=10, leftIndent=28),
    ]
    story += [toc, PageBreak()]

    children = list(doc._element.body)
    toc_index = None
    for index, child in enumerate(children):
        instruction = " ".join((node.text or "") for node in child.findall(".//" + qn("w:instrText")))
        if "TOC" in instruction:
            toc_index = index
            break

    counters = [0, 0, 0]
    for child in children[(toc_index + 1 if toc_index is not None else 0):]:
        if child.tag == qn("w:p"):
            paragraph = Paragraph(child, doc)
            text = paragraph.text.strip()
            style = paragraph.style.name if paragraph.style else ""
            if not text:
                continue
            if text == "Appendix A":
                story.append(PageBreak())
                flowable = RP("Appendix A", styles["AppendixS"])
                flowable._toc_level = 0
                story.append(flowable)
                continue
            if style == "Heading 1":
                counters = [counters[0] + 1, 0, 0]
                story.append(PageBreak())
                flowable = RP(f"{counters[0]}.0 {html.escape(text)}", styles["H1S"])
                flowable._toc_level = 0
                story.append(flowable)
            elif style == "Heading 2":
                counters[1] += 1
                counters[2] = 0
                flowable = RP(f"{counters[0]}.{counters[1]} {html.escape(text)}", styles["H2S"])
                flowable._toc_level = 1
                story.append(flowable)
            elif style == "Heading 3":
                counters[2] += 1
                flowable = RP(f"{counters[0]}.{counters[1]}.{counters[2]} {html.escape(text)}", styles["H3S"])
                flowable._toc_level = 2
                story.append(flowable)
            elif not text.startswith("[[IF:") and not text.startswith("[[END:"):
                story.append(RP(
                    html.escape(text),
                    styles["BulletS"] if style == "List Paragraph" else styles["BodyS"],
                    bulletText="•" if style == "List Paragraph" else None,
                ))
        elif child.tag == qn("w:tbl"):
            source = DocxTable(child, doc)
            if is_signature_table(source):
                story += signature_pdf_flowables(customer, styles)
                continue
            data = [
                [RP(html.escape(cell.text.strip()), styles["BodyS"]) for cell in row.cells]
                for row in source.rows
            ]
            if data:
                table = Table(data, repeatRows=1, hAlign="LEFT")
                table.setStyle(TableStyle([
                    ("GRID", (0, 0), (-1, -1), .45, colors.HexColor("#7d8d94")),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8eef1")),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 7.5),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ]))
                story += [Spacer(1, 5), table, Spacer(1, 5)]

    pdf.multiBuild(story, canvasmaker=_canvas(customer, estimate, _logo(docx_bytes)))
    return out.getvalue()


def install_sow_signature_layout() -> None:
    global _installed
    if _installed:
        return

    previous_mep_docx = sow_service.render_docx
    previous_cip_docx = cip_docx.render_cip_docx

    def render_mep_docx_with_source_signature(db, sow, rev):
        raw = previous_mep_docx(db, sow, rev)
        template = db.get(SOWTemplateVersion, sow.template_version_id)
        if template is not None and template.version_no >= 3:
            return _compact_docx_bytes(raw)
        return raw

    def render_cip_docx_with_source_signature(db, sow, rev):
        return _compact_docx_bytes(previous_cip_docx(db, sow, rev))

    def review_mep_pdf(docx_bytes: bytes, customer: str, estimate: str, sow_date: str) -> bytes:
        return _review_pdf(docx_bytes, customer, estimate, sow_date, "MEP")

    def review_cip_pdf(docx_bytes: bytes, customer: str, estimate: str, sow_date: str) -> bytes:
        return _review_pdf(docx_bytes, customer, estimate, sow_date, "CIP")

    # Word rendering: preserve historical MEP v1/v2 presentation but correct the current MEP
    # template and all CIP SOWs. Blank paragraphs are excluded from canonical text, so the
    # approved-content hash remains stable while the visual layout is corrected.
    sow_service.render_docx = render_mep_docx_with_source_signature
    cip_docx.render_cip_docx = render_cip_docx_with_source_signature
    cip_pdf.render_cip_docx = render_cip_docx_with_source_signature

    # PDF rendering: retain the accepted cover/TOC/header/footer engine and special-case only the
    # signature table so Section 8 matches the two-column format in the supplied source templates.
    sow_pdf_v3._review_pdf = review_mep_pdf
    cip_pdf._review_cip_pdf = review_cip_pdf

    _installed = True
