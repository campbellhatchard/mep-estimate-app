from __future__ import annotations

import base64
import hashlib
import io
import os
import re
import secrets
import zipfile
from dataclasses import dataclass

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from fastapi import Depends, HTTPException, Request
from fastapi.responses import Response
from lxml import etree
from pypdf import PdfReader
from sqlalchemy.orm import Session

from .cip_domain import _take_route
from .cip_models import PRODUCT_CIP, PRODUCT_MEP
from .database import get_db
from .models import EstimateRevision
from .services.audit import record
from .sow_models import SOW, SOWTemplateVersion
from . import sow_service
from .cip_sow import docx as cip_docx
from .cip_sow import pdf as cip_pdf
from .cip_sow.core import SOW_TEMPLATE_CIP_NET_NEW

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
V_NS = "urn:schemas-microsoft-com:vml"
O_NS = "urn:schemas-microsoft-com:office:office"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

NS = {"w": W_NS, "v": V_NS, "o": O_NS, "r": R_NS}
WORD_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
TRACK_PASSWORD_ENV = "SOW_TRACK_CHANGES_PASSWORD"
SPIN_COUNT = 100_000
COVER_FOOTER = (
    "This statement of work estimate is the property and proprietary to Data Systems International, Inc. "
    "dba Cloud inventory® and contains trade secret and confidential information and is solely for Customer’s internal use. "
    "Without the express written consent of Cloud Inventory ®, this estimate shall not be used, reproduced, copied, disclosed, "
    "transmitted in whole or in part. Copyright © 2026 Data Systems International, Inc. dba Cloud Inventory®. All rights reserved."
)


def _w(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"


def _xml(root: etree._Element) -> bytes:
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def _password_verifier(password: str, *, spin_count: int = SPIN_COUNT) -> tuple[str, str]:
    if not password:
        raise ValueError("A non-empty Track Changes password is required.")
    salt = secrets.token_bytes(16)
    value = hashlib.sha512(salt + password.encode("utf-16le")).digest()
    for index in range(spin_count):
        value = hashlib.sha512(value + index.to_bytes(4, byteorder="little")).digest()
    return base64.b64encode(value).decode("ascii"), base64.b64encode(salt).decode("ascii")


@dataclass(frozen=True)
class _Heading:
    level: int
    body_number: str
    toc_number: str
    title: str


def _normalized(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def _heading_entries(doc: Document) -> list[_Heading]:
    counters = [0, 0, 0]
    entries: list[_Heading] = []
    for paragraph in doc.paragraphs:
        style = paragraph.style.name if paragraph.style else ""
        title = paragraph.text.strip()
        if not title:
            continue
        if style == "Heading 1":
            counters = [counters[0] + 1, 0, 0]
            entries.append(_Heading(1, f"{counters[0]}.0", f"{counters[0]}.", title))
        elif style == "Heading 2":
            counters[1] += 1
            counters[2] = 0
            number = f"{counters[0]}.{counters[1]}"
            entries.append(_Heading(2, number, number, title))
        elif style == "Heading 3":
            counters[2] += 1
            number = f"{counters[0]}.{counters[1]}.{counters[2]}"
            entries.append(_Heading(3, number, number, title))
    if not entries:
        raise ValueError("The generated SOW does not contain any numbered headings for its Table of Contents.")
    return entries


def _pdf_layout(pdf_bytes: bytes, headings: list[_Heading]) -> tuple[dict[tuple[str, str], int], dict[tuple[str, str], int]]:
    reader = PdfReader(io.BytesIO(pdf_bytes))
    pages = [_normalized(page.extract_text() or "") for page in reader.pages]
    if not pages:
        raise ValueError("The SOW review PDF does not contain any pages.")

    first = headings[0]
    first_needle = _normalized(f"{first.body_number} {first.title}")
    # The TOC page(s) list every heading using the same numbered text as the heading's own
    # body page (by construction, both the TOC entry and the heading text come from the same
    # "{number} {title}" string). Scanning forward and taking the FIRST match therefore finds
    # the TOC listing, not the real body page, because the TOC always precedes the body.
    # Take the LAST match instead: the true body occurrence is always the latest one, since
    # every heading is preceded by a page break and the TOC is fully emitted before any body
    # content begins.
    first_matches = [index + 1 for index, text in enumerate(pages) if first_needle in text]
    body_start = first_matches[-1] if first_matches else None
    if body_start is None:
        raise ValueError("The SOW review PDF could not be reconciled to the Word document headings.")

    body_pages: dict[tuple[str, str], int] = {}
    toc_pages: dict[tuple[str, str], int] = {}
    for heading in headings:
        key = (heading.body_number, heading.title)
        body_needle = _normalized(f"{heading.body_number} {heading.title}")
        body_page = next(
            (index + 1 for index in range(body_start - 1, len(pages)) if body_needle in pages[index]),
            None,
        )
        if body_page is None:
            raise ValueError(f"The review PDF is missing the heading '{heading.title}'.")
        body_pages[key] = body_page

        toc_needles = (
            body_needle,
            _normalized(f"{heading.toc_number} {heading.title}"),
        )
        toc_page = None
        for index in range(1, max(body_start - 1, 1)):
            text = pages[index]
            if any(needle and needle in text for needle in toc_needles):
                toc_page = index + 1
                break
        if toc_page is None:
            toc_page = 2
        toc_pages[key] = toc_page

    current = min(toc_pages.values()) if toc_pages else 2
    for heading in headings:
        key = (heading.body_number, heading.title)
        value = toc_pages[key]
        if value < current:
            toc_pages[key] = current
        else:
            current = value
    return body_pages, toc_pages


def _extract_logo(docx_bytes: bytes) -> bytes:
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as archive:
        images = [
            name for name in archive.namelist()
            if name.startswith("word/media/") and name.lower().endswith((".png", ".jpg", ".jpeg"))
        ]
        if not images:
            raise ValueError("The SOW template does not contain the Cloud Inventory logo.")
        return archive.read(images[0])


def _clear_part(container) -> None:
    root = container._element
    for child in list(root):
        root.remove(child)
    root.append(OxmlElement("w:p"))


def _set_paragraph_bottom_border(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    borders = ppr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        ppr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")


def _set_table_bottom_border(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")


def _add_field(paragraph, instruction: str, placeholder: str = "1") -> None:
    run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    code = OxmlElement("w:instrText")
    code.set(qn("xml:space"), "preserve")
    code.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, code, separate, text, end):
        run.append(node)
    paragraph._p.append(run)


def _watermark_paragraph(text: str = "DRAFT"):
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    pict = OxmlElement("w:pict")
    shape = etree.Element(f"{{{V_NS}}}shape", nsmap={"v": V_NS, "o": O_NS})
    shape.set("id", "PowerPlusWaterMarkObject357476642")
    shape.set(f"{{{O_NS}}}spid", "_x0000_s2049")
    shape.set("type", "#_x0000_t136")
    shape.set(
        "style",
        "position:absolute;margin-left:0;margin-top:0;width:527.85pt;height:131.95pt;"
        "rotation:315;z-index:-251654144;mso-position-horizontal:center;"
        "mso-position-horizontal-relative:margin;mso-position-vertical:center;"
        "mso-position-vertical-relative:margin;mso-wrap-edited:f",
    )
    shape.set("fillcolor", "#C0C0C0")
    shape.set("stroked", "f")
    fill = etree.SubElement(shape, f"{{{V_NS}}}fill")
    fill.set("opacity", "0.5")
    textpath = etree.SubElement(shape, f"{{{V_NS}}}textpath")
    textpath.set("style", 'font-family:"Calibri";font-size:1pt')
    textpath.set("string", text)
    path = etree.SubElement(shape, f"{{{V_NS}}}path")
    path.set("textpathok", "t")
    lock = etree.SubElement(shape, f"{{{O_NS}}}lock")
    lock.set(f"{{{V_NS}}}ext", "edit")
    lock.set("aspectratio", "t")
    pict.append(shape)
    run.append(pict)
    paragraph.append(run)
    return paragraph


def _build_headers_and_footers(doc: Document, logo_bytes: bytes, customer: str, estimate: str, *, draft: bool) -> None:
    for section in doc.sections:
        section.different_first_page_header_footer = True
        section.header_distance = Inches(0.25)
        section.footer_distance = Inches(0.25)
        for part in (section.header, section.first_page_header, section.footer, section.first_page_footer):
            part.is_linked_to_previous = False
            _clear_part(part)

        first_header = section.first_page_header
        first_header_paragraph = first_header.paragraphs[0]
        first_header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        first_header_paragraph.add_run().add_picture(io.BytesIO(logo_bytes), width=Inches(1.45))
        _set_paragraph_bottom_border(first_header_paragraph)

        header = section.header
        _clear_part(header)
        header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        left = header_table.cell(0, 0).paragraphs[0]
        left.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left.add_run().add_picture(io.BytesIO(logo_bytes), width=Inches(1.45))
        right = header_table.cell(0, 1).paragraphs[0]
        right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        customer_run = right.add_run(customer)
        customer_run.bold = True
        customer_run.font.size = Pt(9)
        _set_table_bottom_border(header_table)

        cover_footer = section.first_page_footer.paragraphs[0]
        cover_footer.text = COVER_FOOTER
        cover_footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in cover_footer.runs:
            run.font.name = "Arial"
            run.font.size = Pt(6.3)

        footer = section.footer
        _clear_part(footer)
        footer_table = footer.add_table(rows=1, cols=3, width=Inches(6.5))
        footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        estimate_paragraph = footer_table.cell(0, 0).paragraphs[0]
        estimate_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        estimate_run = estimate_paragraph.add_run(f"Estimate Number: {estimate}")
        estimate_run.bold = True
        estimate_run.font.size = Pt(7.5)
        page_paragraph = footer_table.cell(0, 1).paragraphs[0]
        page_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        page_run = page_paragraph.add_run("Page ")
        page_run.bold = True
        page_run.font.size = Pt(7.5)
        _add_field(page_paragraph, "PAGE")
        of_run = page_paragraph.add_run(" of ")
        of_run.bold = True
        of_run.font.size = Pt(7.5)
        _add_field(page_paragraph, "NUMPAGES")

        if draft:
            first_header._element.append(_watermark_paragraph())
            header._element.append(_watermark_paragraph())


def _toc_paragraph(heading: _Heading, page: int, *, page_break_before: bool):
    paragraph = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    paragraph.append(ppr)
    if page_break_before:
        ppr.append(OxmlElement("w:pageBreakBefore"))
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), str({1: 240, 2: 220, 3: 200}[heading.level]))
    spacing.set(qn("w:lineRule"), "exact")
    ppr.append(spacing)
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), "9000")
    tabs.append(tab)
    ppr.append(tabs)
    if heading.level > 1:
        indent = OxmlElement("w:ind")
        indent.set(qn("w:left"), str((heading.level - 1) * 280))
        ppr.append(indent)

    def add_text(value: str, *, bold: bool = False) -> None:
        run = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        if bold:
            rpr.append(OxmlElement("w:b"))
        size = OxmlElement("w:sz")
        size.set(qn("w:val"), str({1: 18, 2: 17, 3: 16}[heading.level]))
        rpr.append(size)
        run.append(rpr)
        text = OxmlElement("w:t")
        text.text = value
        run.append(text)
        paragraph.append(run)

    add_text(f"{heading.toc_number} {heading.title}", bold=heading.level == 1)
    tab_run = OxmlElement("w:r")
    tab_run.append(OxmlElement("w:tab"))
    paragraph.append(tab_run)
    add_text(str(page), bold=heading.level == 1)
    return paragraph


def _replace_toc_with_static_pages(doc: Document, headings: list[_Heading], body_pages, toc_pages) -> None:
    body = doc._element.body
    toc = None
    for child in list(body):
        instruction = " ".join(node.text or "" for node in child.findall(".//" + qn("w:instrText")))
        if "TOC" in instruction:
            toc = child
            break
    if toc is None:
        raise ValueError("The generated SOW does not contain a Word Table of Contents field.")

    title = OxmlElement("w:p")
    title_ppr = OxmlElement("w:pPr")
    title.append(title_ppr)
    center = OxmlElement("w:jc")
    center.set(qn("w:val"), "center")
    title_ppr.append(center)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "240")
    title_ppr.append(spacing)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rpr.append(OxmlElement("w:b"))
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "32")
    rpr.append(size)
    run.append(rpr)
    text = OxmlElement("w:t")
    text.text = "Table of Contents"
    run.append(text)
    title.append(run)
    toc.addprevious(title)

    previous_toc_page = None
    for heading in headings:
        key = (heading.body_number, heading.title)
        toc_page = toc_pages[key]
        entry = _toc_paragraph(
            heading,
            body_pages[key],
            page_break_before=(previous_toc_page is not None and toc_page > previous_toc_page),
        )
        toc.addprevious(entry)
        previous_toc_page = max(previous_toc_page or toc_page, toc_page)
    body.remove(toc)


def _remove_automatic_field_refresh(doc: Document) -> None:
    settings = doc.settings._element
    for node in list(settings.findall(qn("w:updateFields"))):
        settings.remove(node)
    roots = [doc._element]
    for section in doc.sections:
        roots.extend((
            section.header._element,
            section.first_page_header._element,
            section.footer._element,
            section.first_page_footer._element,
        ))
    for root in roots:
        for field in root.findall(".//" + qn("w:fldChar")):
            field.attrib.pop(qn("w:dirty"), None)


def prepare_word_presentation(
    docx_bytes: bytes,
    *,
    pdf_bytes: bytes,
    customer: str,
    estimate: str,
    draft: bool,
) -> bytes:
    try:
        doc = Document(io.BytesIO(docx_bytes))
    except Exception as exc:
        raise ValueError("The generated SOW is not a valid Word .docx document.") from exc

    if doc._element.findall(".//" + qn("w:ins")) or doc._element.findall(".//" + qn("w:del")):
        raise ValueError("The generated SOW unexpectedly contains tracked revisions before download.")

    headings = _heading_entries(doc)
    body_pages, toc_pages = _pdf_layout(pdf_bytes, headings)
    _replace_toc_with_static_pages(doc, headings, body_pages, toc_pages)
    _build_headers_and_footers(doc, _extract_logo(docx_bytes), customer, estimate, draft=draft)
    _remove_automatic_field_refresh(doc)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def _settings_with_protection(settings_xml: bytes, password: str) -> bytes:
    root = etree.fromstring(settings_xml)
    for node in root.findall("w:trackRevisions", namespaces=NS):
        root.remove(node)
    root.insert(0, etree.Element(_w("trackRevisions")))
    for node in root.findall("w:documentProtection", namespaces=NS):
        root.remove(node)
    for node in root.findall("w:updateFields", namespaces=NS):
        root.remove(node)

    hash_value, salt_value = _password_verifier(password)
    protection = etree.Element(
        _w("documentProtection"),
        {
            _w("edit"): "trackedChanges",
            _w("enforcement"): "1",
            _w("algorithmName"): "SHA-512",
            _w("hashValue"): hash_value,
            _w("saltValue"): salt_value,
            _w("spinCount"): str(SPIN_COUNT),
        },
    )
    root.insert(1 if len(root) else 0, protection)
    return _xml(root)


def apply_word_controls(docx_bytes: bytes, *, password: str, draft: bool) -> bytes:
    source = io.BytesIO(docx_bytes)
    output = io.BytesIO()
    try:
        archive = zipfile.ZipFile(source, "r")
    except zipfile.BadZipFile as exc:
        raise ValueError("The generated SOW is not a valid Word .docx document.") from exc

    with archive:
        names = archive.namelist()
        if "word/settings.xml" not in names or "word/document.xml" not in names:
            raise ValueError("The Word document is missing required Open XML parts.")
        overrides = {
            "word/settings.xml": _settings_with_protection(archive.read("word/settings.xml"), password)
        }
        with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as destination:
            for info in archive.infolist():
                destination.writestr(info, overrides.get(info.filename, archive.read(info.filename)))

    controlled = output.getvalue()
    with zipfile.ZipFile(io.BytesIO(controlled), "r") as archive:
        settings = etree.fromstring(archive.read("word/settings.xml"))
        document = etree.fromstring(archive.read("word/document.xml"))
        protection = settings.find("w:documentProtection", namespaces=NS)
        track = settings.find("w:trackRevisions", namespaces=NS)
        if (
            protection is None
            or track is None
            or protection.get(_w("edit")) != "trackedChanges"
            or protection.get(_w("enforcement")) != "1"
            or protection.get(_w("algorithmName")) != "SHA-512"
            or not protection.get(_w("hashValue"))
            or not protection.get(_w("saltValue"))
        ):
            raise ValueError("Controlled Word protection could not be verified.")
        if settings.find("w:updateFields", namespaces=NS) is not None:
            raise ValueError("Controlled Word field refresh must be disabled before Track Changes is enabled.")
        if document.findall(".//w:ins", namespaces=NS) or document.findall(".//w:del", namespaces=NS):
            raise ValueError("Controlled Word contains tracked revisions before the user has edited it.")
        if any(field.get(_w("dirty")) for field in document.findall(".//w:fldChar", namespaces=NS)):
            raise ValueError("Controlled Word contains dirty body fields that could create initial tracked revisions.")

        section = document.find(".//w:sectPr", namespaces=NS)
        if section is None or section.find("w:titlePg", namespaces=NS) is None:
            raise ValueError("Controlled Word first-page header/footer configuration could not be verified.")
        header_refs = section.findall("w:headerReference", namespaces=NS)
        footer_refs = section.findall("w:footerReference", namespaces=NS)
        header_types = {ref.get(_w("type")) for ref in header_refs}
        footer_types = {ref.get(_w("type")) for ref in footer_refs}
        if not {"first", "default"}.issubset(header_types) or not {"first", "default"}.issubset(footer_types):
            raise ValueError("Controlled Word header/footer relationships could not be verified.")

        header_parts = [archive.read(name) for name in names if re.fullmatch(r"word/header\d+\.xml", name)]
        footer_parts = [archive.read(name) for name in names if re.fullmatch(r"word/footer\d+\.xml", name)]
        watermark_count = sum(raw.count(b"PowerPlusWaterMarkObject") for raw in header_parts)
        if draft and watermark_count < 2:
            raise ValueError("DRAFT Printed Watermark could not be verified on the controlled Word document.")
        if not draft and watermark_count:
            raise ValueError("Approved Word document unexpectedly contains a DRAFT watermark.")
        if not any(COVER_FOOTER.encode("utf-8")[:40] in raw for raw in footer_parts):
            raise ValueError("Controlled Word cover-page proprietary footer could not be verified.")
        if not any(b"Estimate Number:" in raw and b"PAGE" in raw and b"NUMPAGES" in raw for raw in footer_parts):
            raise ValueError("Controlled Word page footer could not be verified.")

    if password.encode("utf-8") in controlled:
        raise ValueError("Controlled Word protection failed because clear-text secret data was detected.")
    return controlled


def _raw_docx_for_sow(db: Session, sow: SOW, rev: EstimateRevision) -> bytes:
    template = db.get(SOWTemplateVersion, sow.template_version_id)
    if not template:
        raise ValueError("The SOW template version no longer exists.")
    if template.template_key == "MEP_NET_NEW":
        return (
            sow_service.verify_approved_content(db, sow, rev)
            if sow.status == "APPROVED"
            else sow_service.render_docx(db, sow, rev)
        )
    if template.template_key == SOW_TEMPLATE_CIP_NET_NEW:
        return (
            cip_docx.verify_cip_approved_content(db, sow, rev)
            if sow.status == "APPROVED"
            else cip_docx.render_cip_docx(db, sow, rev)
        )
    raise ValueError("Controlled Word download is not available for this SOW template family.")


def _review_pdf_for_sow(db: Session, sow: SOW, rev: EstimateRevision) -> bytes:
    template = db.get(SOWTemplateVersion, sow.template_version_id)
    if not template:
        raise ValueError("The SOW template version no longer exists.")
    if template.template_key == "MEP_NET_NEW":
        return sow_service.render_pdf(db, sow, rev)
    if template.template_key == SOW_TEMPLATE_CIP_NET_NEW:
        return cip_pdf.render_cip_pdf(db, sow, rev)
    raise ValueError("SOW review PDF is not available for this SOW template family.")


def _word_filename(sow: SOW, rev: EstimateRevision, product: str) -> str:
    status_suffix = "" if sow.status == "APPROVED" else "-DRAFT"
    product_segment = "CIP" if product == PRODUCT_CIP else "MEP"
    return f"{rev.estimate.estimate_number}-{product_segment}-SOW-R{sow.sow_revision_no}{status_suffix}.docx"


def register_controlled_sow_word(app, core) -> None:
    _take_route(app, "/sow/{sid}/docx", "GET")

    @app.get("/sow/{sid}/docx")
    def controlled_sow_docx(sid: int, request: Request, db: Session = Depends(get_db)):
        user = core.current_user(request, db)
        sow = db.get(SOW, sid)
        if not sow:
            raise HTTPException(404, "SOW not found")
        rev = db.get(EstimateRevision, sow.estimate_revision_id)
        if not rev:
            raise HTTPException(404, "Estimate revision not found")

        password = os.getenv(TRACK_PASSWORD_ENV, "").strip()
        if not password:
            raise HTTPException(
                503,
                "Controlled Word download is unavailable because SOW_TRACK_CHANGES_PASSWORD is not configured.",
            )

        try:
            raw = _raw_docx_for_sow(db, sow, rev)
            review_pdf = _review_pdf_for_sow(db, sow, rev)
            presented = prepare_word_presentation(
                raw,
                pdf_bytes=review_pdf,
                customer=rev.customer or "",
                estimate=rev.estimate.estimate_number,
                draft=sow.status != "APPROVED",
            )
            content = apply_word_controls(
                presented,
                draft=sow.status != "APPROVED",
                password=password,
            )
        except ValueError as exc:
            raise HTTPException(409, str(exc)) from exc

        event_type = "SOW_APPROVED_DOCX_GENERATED" if sow.status == "APPROVED" else "SOW_DRAFT_DOCX_GENERATED"
        product = PRODUCT_CIP if (rev.engine_version or "").upper().startswith("CIP-") else PRODUCT_MEP
        record(
            db,
            event_type=event_type,
            user_id=user.id,
            estimate_id=rev.estimate_id,
            revision_id=rev.id,
            field_name=f"SOW:{sow.id}",
            new_value=f"Template {sow.template_version_id}; content {sow.content_hash or 'unapproved'}",
        )
        db.commit()

        return Response(
            content,
            media_type=WORD_MIME,
            headers={"Content-Disposition": f'attachment; filename="{_word_filename(sow, rev, product)}"'},
        )
