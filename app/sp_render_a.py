from __future__ import annotations

import re
from copy import deepcopy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from sqlalchemy.orm import Session

from . import sow_service
from .cip_models import CIPRevisionInput, PRODUCT_CIP
from .models import EstimateRevision
from .sow_models import SOW
from .small_project_models import SmallProjectSOWConfig
from .sp_core_a import (
    METHODOLOGY_SPECS, SMALL_PROJECT_TEMPLATE_KEYS, WEEKEND_HOLIDAY_CLAUSE,
    _config, _product_for_revision, _template_for_sow, small_project_support_hours,
)
from .sp_core_b import _deliverable_map, _methodology_map, methodology_included, appendix_included


def _paragraph_text(el) -> str:
    return "".join(t.text or "" for t in el.findall(".//" + qn("w:t"))).strip()


def _paragraph_style(el) -> str:
    ppr = el.find(qn("w:pPr"))
    if ppr is None:
        return ""
    style = ppr.find(qn("w:pStyle"))
    return style.get(qn("w:val"), "") if style is not None else ""


def _delete_element(el) -> None:
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def _set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _insert_paragraph_before(target: Paragraph, text: str, style: str = "Normal") -> Paragraph:
    el = OxmlElement("w:p")
    target._p.addprevious(el)
    paragraph = Paragraph(el, target._parent)
    paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def _replace_everywhere(doc: Document, replacements: dict[str, str]) -> None:
    for paragraph in list(sow_service._all_doc_paragraphs(doc)):
        text = paragraph.text
        if not any(token in text for token in replacements):
            continue
        for token, value in replacements.items():
            text = text.replace(token, value)
        _set_paragraph_text(paragraph, text)


def _replace_project_objective(doc: Document, objective: str) -> None:
    paragraphs = list(doc.paragraphs)
    heading_index = next(
        (i for i, p in enumerate(paragraphs) if p.text.strip() == "Project Objective"),
        None,
    )
    deliverables_index = next(
        (i for i, p in enumerate(paragraphs) if p.text.strip() == "Deliverables"),
        None,
    )
    if heading_index is None or deliverables_index is None:
        raise ValueError("Small Project template is missing Project Objective or Deliverables.")
    body = paragraphs[heading_index + 1:deliverables_index]
    if body:
        _set_paragraph_text(body[0], objective)
        body[0].style = "Normal"
        for paragraph in body[1:]:
            _delete_element(paragraph._p)
    else:
        _insert_paragraph_before(paragraphs[deliverables_index], objective, "Normal")


def _replace_deliverables(doc: Document, cfg: SmallProjectSOWConfig) -> None:
    paragraphs = list(doc.paragraphs)
    software = next(
        (p for p in paragraphs if p.text.strip() == "Software Deliverables"), None
    )
    methodology = next(
        (p for p in paragraphs if p.text.strip() == "Service Methodology"), None
    )
    if not software or not methodology:
        raise ValueError("Small Project template is missing Software Deliverables or Service Methodology.")

    body = doc._element.body
    children = list(body)
    start = children.index(software._p) + 1
    end = children.index(methodology._p)
    for el in children[start:end]:
        _delete_element(el)

    for row in sorted(cfg.deliverables, key=lambda x: x.sort_order):
        if not row.include:
            continue
        _insert_paragraph_before(methodology, row.name, "Heading 3")
        if row.scope_description.strip():
            _insert_paragraph_before(methodology, row.scope_description.strip(), "List Paragraph")
        for line in re.split(r"[\r\n]+", row.detail_notes or ""):
            if line.strip():
                _insert_paragraph_before(methodology, line.strip(), "List Paragraph")


def _methodology_blocks(doc: Document) -> dict[str, list]:
    body = doc._element.body
    children = list(body)
    service = next(
        (el for el in children if el.tag == qn("w:p") and _paragraph_text(el) == "Service Methodology"),
        None,
    )
    assumptions = next(
        (el for el in children if el.tag == qn("w:p") and _paragraph_text(el) == "Solution Assumptions"),
        None,
    )
    if service is None or assumptions is None:
        raise ValueError("Small Project template is missing Service Methodology or Solution Assumptions.")
    start = children.index(service)
    end = children.index(assumptions)
    blocks: dict[str, list] = {}
    current_title = None
    for el in children[start + 1:end]:
        if el.tag == qn("w:p") and _paragraph_style(el) == "Heading3":
            current_title = _paragraph_text(el)
            blocks[current_title] = [deepcopy(el)]
        elif current_title is not None:
            blocks[current_title].append(deepcopy(el))
    return blocks


def _filter_methodology(
    doc: Document,
    db: Session,
    sow: SOW,
    rev: EstimateRevision,
    cfg: SmallProjectSOWConfig,
) -> None:
    blocks = _methodology_blocks(doc)
    body = doc._element.body
    children = list(body)
    project_structure = next(
        (
            el for el in children
            if el.tag == qn("w:p") and _paragraph_text(el) == "Project Structure"
        ),
        None,
    )
    assumptions = next(
        (
            el for el in children
            if el.tag == qn("w:p") and _paragraph_text(el) == "Solution Assumptions"
        ),
        None,
    )
    if project_structure is None or assumptions is None:
        raise ValueError("Small Project template is missing Project Structure or Solution Assumptions.")

    children = list(body)
    start = children.index(project_structure) + 1
    end = children.index(assumptions)
    for el in children[start:end]:
        _delete_element(el)

    modes = _methodology_map(cfg)
    for key, title in METHODOLOGY_SPECS:
        row = modes.get(key)
        if not row or not methodology_included(db, sow, rev, cfg, row):
            continue
        for el in blocks.get(title, []):
            assumptions.addprevious(deepcopy(el))


def _delete_table(table) -> None:
    _delete_element(table._element)


def _delete_row(table, row) -> None:
    table._tbl.remove(row._tr)


def _fill_hypercare_table(doc: Document, sow: SOW, support_hours: float) -> None:
    table = next(
        (
            t for t in doc.tables
            if t.rows and t.rows[0].cells[0].text.strip() == "Location Description"
        ),
        None,
    )
    if not table:
        return
    while len(table.rows) > 1:
        _delete_row(table, table.rows[-1])
    active = [
        row for row in sow.hypercare_locations
        if float(row.allocated_hours or 0) > 0 or row.description.strip()
    ]
    if support_hours <= 0 or not active:
        _delete_table(table)
        return
    for row in active:
        cells = table.add_row().cells
        cells[0].text = row.description
        cells[1].text = row.support_type
        cells[2].text = row.country
        cells[3].text = f"{float(row.allocated_hours or 0):g}"


def _fill_commercial_table(doc: Document, rev: EstimateRevision) -> None:
    table = next(
        (
            t for t in doc.tables
            if t.rows and "Approved Hourly Rate" in t.rows[0].cells[0].text
        ),
        None,
    )
    if table and len(table.rows) > 1:
        table.rows[1].cells[0].text = f"{rev.billing_rate:,.2f}"
        table.rows[1].cells[1].text = f"{rev.calculated_hours:g}"
        table.rows[1].cells[2].text = f"{rev.calculated_fees:,.2f}"


def _fill_appendix(
    doc: Document,
    db: Session,
    sow: SOW,
    rev: EstimateRevision,
    cfg: SmallProjectSOWConfig,
    product: str,
) -> None:
    table = next(
        (
            t for t in doc.tables
            if t.rows and t.rows[0].cells[0].text.strip() == "Deployment Point"
        ),
        None,
    )
    appendix = next((p for p in doc.paragraphs if p.text.strip() == "Appendix A"), None)
    planned = next((p for p in doc.paragraphs if p.text.strip() == "Planned Deployment Summary"), None)
    intro = next(
        (
            p for p in doc.paragraphs
            if p.text.strip().startswith("This section defines the initial deployment details")
        ),
        None,
    )
    if not appendix_included(db, sow, rev):
        if table:
            _delete_table(table)
        for p in (appendix, planned, intro):
            if p:
                _delete_element(p._p)
        return
    if not table:
        raise ValueError("Small Project template is missing Appendix A deployment table.")

    inp = db.get(CIPRevisionInput, rev.id) if product == PRODUCT_CIP else None
    epp = _deliverable_map(cfg).get("EPP")
    epp_included = bool(epp and epp.include)
    erp_solution = inp.deployed_over if inp else rev.erp
    deployment_model = (
        "Cloud Inventory® Managed / Public Cloud"
        if product == PRODUCT_CIP or cfg.install_mode == "Cloud"
        else "Customer Managed / On Premises"
    )

    values = {
        "Planned ERP Solution": f"{erp_solution} {sow.erp_version}".strip(),
        "Planned ERP Base Code Version": sow.erp_base_code_version,
        "Planned ERP Tools Release Version": sow.erp_tools_release,
        "Planned ERP Operating System Version (App Server)": sow.erp_os_version,
        "Planned ERP Database Type / Version": sow.erp_database_version,
        "Planned Cloud Inventory® Product / Version": sow.mep_product_version,
        "Planned Label Printing Solution": sow.epp_product_version if epp_included else "",
        "Planned Print Method(s)": sow.print_methods if epp_included else "",
        "Planned Deployment Model – Cloud Inventory® Product": deployment_model,
        "Planned Deployment Model – ERP Solution": sow.erp_deployment_model,
    }

    for row in list(table.rows)[1:]:
        left = row.cells[0].text.strip() if row.cells else ""
        if left == "Solution Deployment":
            continue
        if left.startswith("Planned Device Information") or left in (
            "Handheld Units", "Vehicle Mount Units", "Desktop Environment", ""
        ):
            _delete_row(table, row)
            continue
        if left in values:
            value = values[left]
            if not value and left not in ("Planned ERP Solution",):
                _delete_row(table, row)
            else:
                row.cells[1].text = value


def _replace_commercial_wording(doc: Document, rev: EstimateRevision, sow: SOW) -> None:
    for p in list(doc.paragraphs):
        text = p.text.strip()
        if text.startswith("Costs shown herein"):
            _set_paragraph_text(p, f"Costs shown herein are in {rev.currency} currency.")
        elif text.startswith("Approved Hourly Rate stated above"):
            _set_paragraph_text(
                p,
                "Approved Hourly Rate stated above is for work performed during standard business weekdays.",
            )
            next_p = p._p.getnext()
            next_text = _paragraph_text(next_p) if next_p is not None and next_p.tag == qn("w:p") else ""
            if next_text != WEEKEND_HOLIDAY_CLAUSE:
                inserted = OxmlElement("w:p")
                p._p.addnext(inserted)
                pp = Paragraph(inserted, p._parent)
                pp.style = "List Paragraph"
                pp.add_run(WEEKEND_HOLIDAY_CLAUSE)
        elif text.startswith("Cloud Inventory ® will submit invoices to Customer on a weekly basis"):
            _set_paragraph_text(
                p,
                text.replace(
                    "on a weekly basis",
                    f"on a {sow.invoice_frequency.lower()} basis",
                    1,
                ),
            )


def _replace_agreement(doc: Document, sow: SOW) -> None:
    for p in doc.paragraphs:
        if p.text.strip().startswith("Unless stated otherwise herein, this SOW is governed"):
            _set_paragraph_text(
                p,
                "Unless stated otherwise herein, this SOW is governed by the terms and conditions of the "
                f"{sow.agreement_type} between the parties (the “Agreement”). Any capitalized terms in this "
                "SOW not defined herein shall have the meaning assigned to them in the Agreement.",
            )
            break


def _replace_key_user_count(doc: Document, count: int) -> None:
    replacements = {
        "up to 5 users": f"up to {count} user{'s' if count != 1 else ''}",
        "up to two (2) key users": f"up to {count} key user{'s' if count != 1 else ''}",
    }
    for p in list(sow_service._all_doc_paragraphs(doc)):
        text = p.text
        updated = text
        for old, new in replacements.items():
            updated = updated.replace(old, new)
        if updated != text:
            _set_paragraph_text(p, updated)
