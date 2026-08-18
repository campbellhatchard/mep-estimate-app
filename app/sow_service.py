from __future__ import annotations

import base64
import hashlib
import io
import html
from copy import deepcopy
from datetime import date, datetime
from pathlib import Path
from zipfile import BadZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table as DocxTable
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph as RLParagraph, SimpleDocTemplate, Spacer, Table, TableStyle
from sqlalchemy import desc
from sqlalchemy.orm import Session

from .assumptions import EstimateAssumption
from .cip_models import PRODUCT_MEP
from .models import EstimateApplication, EstimateCustomApplication, EstimateRevision, User
from .services.audit import record
from .services.calculation import calculation
from .sow_models import SOW, SOWDevice, SOWHypercareLocation, SOWTemplateVersion, SOW_TEMPLATE_MEP_NET_NEW

BASE = Path(__file__).parent
INITIAL_TEMPLATE_PART_GLOB = "mep_new_client_v1.b64.part*"
INITIAL_TEMPLATE_DIR = BASE / "sow_templates"

AGREEMENT_TYPES = (
    "Software as a Service Agreement",
    "Software License, Services and Maintenance Agreement",
)
INVOICE_FREQUENCIES = ("Weekly", "Monthly")
SUPPORT_TYPES = ("Remote", "On-Site")
DEVICE_TYPES = ("Handheld Unit", "Vehicle Mount Unit", "Desktop Environment", "Other")

REQUIRED_TEMPLATE_MARKERS = {
    "[[CUSTOMER_NAME]]", "[[ESTIMATE_NUMBER]]", "[[SOW_DATE]]", "[[ENTITY]]",
    "[[AGREEMENT_TYPE]]", "[[PROJECT_OBJECTIVE]]", "[[EPP_INSTALL_MODE]]",
    "[[BARCODE_PRINTER_COUNT]]", "[[LABEL_COUNT]]", "[[BASELINE_APPLICATIONS]]",
    "[[CUSTOM_APPLICATIONS]]", "[[LABELS]]", "[[ASSUMPTIONS]]",
    "[[BILLING_RATE]]", "[[ESTIMATED_HOURS]]", "[[ESTIMATED_COST]]",
    "[[INVOICE_FREQUENCY]]", "[[ERP_SOLUTION]]", "[[ERP_VERSION]]",
    "[[ERP_BASE_CODE]]", "[[ERP_TOOLS_RELEASE]]", "[[ERP_OS]]", "[[ERP_DB]]",
    "[[MEP_VERSION]]", "[[EPP_VERSION]]", "[[PRINT_METHODS]]",
    "[[CI_DEPLOYMENT_MODEL]]", "[[ERP_DEPLOYMENT_MODEL]]", "[[DEVICE_ROWS]]",
    "[[HYPERCARE_ROWS]]", "[[CURRENCY]]", "[[IF:ASSUMPTIONS]]", "[[END:ASSUMPTIONS]]",
}


def sha256_bytes(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()


def template_text(content: bytes) -> str:
    try:
        doc = Document(io.BytesIO(content))
    except (BadZipFile, ValueError, KeyError) as exc:
        raise ValueError("The uploaded file is not a valid Word .docx document.") from exc
    chunks: list[str] = []
    chunks.extend(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    for section in doc.sections:
        chunks.extend(p.text for p in section.header.paragraphs)
        chunks.extend(p.text for p in section.footer.paragraphs)
    return "\n".join(chunks)


def validate_template(content: bytes) -> list[str]:
    text = template_text(content)
    return sorted(marker for marker in REQUIRED_TEMPLATE_MARKERS if marker not in text)


def seed_initial_sow_template(db: Session) -> None:
    if db.query(SOWTemplateVersion).filter(SOWTemplateVersion.template_key == SOW_TEMPLATE_MEP_NET_NEW).count():
        return
    parts = sorted(INITIAL_TEMPLATE_DIR.glob(INITIAL_TEMPLATE_PART_GLOB))
    if not parts:
        return
    admin = db.query(User).filter(User.username_normalized == "admin").first()
    if not admin:
        return
    encoded = "".join(part.read_text() for part in parts)
    content = base64.b64decode(encoded.strip())
    missing = validate_template(content)
    if missing:
        raise RuntimeError(f"Bundled SOW template is missing markers: {', '.join(missing)}")
    row = SOWTemplateVersion(
        template_key=SOW_TEMPLATE_MEP_NET_NEW,
        label="MEP New Client SOW",
        product_type=PRODUCT_MEP,
        customer_type="Net_New",
        version_no=1,
        status="ACTIVE",
        filename="MEP_Template_NewClient_2026_14_Controlled_v1.docx",
        content=content,
        content_sha256=sha256_bytes(content),
        change_reason="Initial controlled MEP New Client SOW template based on MEP_Template_NewClient_2026_14.",
        created_by=admin.id,
        activated_by=admin.id,
        activated_at=datetime.utcnow(),
    )
    db.add(row)
    db.flush()
    record(db, event_type="SOW_TEMPLATE_ACTIVATED", user_id=admin.id,
           field_name=f"SOW_TEMPLATE:{SOW_TEMPLATE_MEP_NET_NEW}:1",
           new_value=row.filename, reason=row.change_reason)
    db.commit()


def active_template(db: Session, template_key: str = SOW_TEMPLATE_MEP_NET_NEW) -> SOWTemplateVersion:
    row = (
        db.query(SOWTemplateVersion)
        .filter(SOWTemplateVersion.template_key == template_key, SOWTemplateVersion.status == "ACTIVE")
        .order_by(desc(SOWTemplateVersion.version_no))
        .first()
    )
    if not row:
        raise ValueError("No active SOW template is available.")
    return row


def sow_eligible(rev: EstimateRevision) -> bool:
    return (
        rev.status in ("APPROVED", "FINAL", "SUPERSEDED")
        and not (rev.engine_version or "").upper().startswith("CIP-")
        and rev.customer_type == "Net_New"
    )


def latest_sow(db: Session, estimate_revision_id: int) -> SOW | None:
    return (
        db.query(SOW)
        .filter(SOW.estimate_revision_id == estimate_revision_id)
        .order_by(desc(SOW.sow_revision_no))
        .first()
    )


def go_live_support_hours(db: Session, rev: EstimateRevision) -> float:
    lines, _, _, _ = calculation(db, rev)
    for line in lines:
        if getattr(line, "key", "") == "GO_LIVE_SUPPORT":
            return float(line.extended_hours or 0)
    return 0.0


def ci_deployment_model(rev: EstimateRevision) -> str:
    if rev.project_type in ("MEP Cloud", "Platform Move To Cloud", "EPP Cloud"):
        return "Cloud Inventory® Managed / Public Cloud"
    if rev.project_type in ("MEP On Prem", "Platform Move On Prem", "EPP On Prem"):
        return "Customer Managed / On Premises"
    return "To be confirmed during Architecture Design Workshop"


def selected_baseline_apps(db: Session, rev: EstimateRevision) -> list[str]:
    return [
        row.label for row in db.query(EstimateApplication)
        .filter(EstimateApplication.revision_id == rev.id, EstimateApplication.kind == "APPLICATION")
        .order_by(EstimateApplication.sort_order).all()
        if row.config_type != "No Config"
    ]


def selected_custom_apps(db: Session, rev: EstimateRevision) -> list[str]:
    return [
        row.description.strip() for row in db.query(EstimateCustomApplication)
        .filter(EstimateCustomApplication.revision_id == rev.id)
        .order_by(EstimateCustomApplication.sort_order).all()
        if row.description.strip() and row.complexity != "No Config"
    ]


def estimate_assumptions(db: Session, rev: EstimateRevision) -> list[str]:
    return [
        row.text.strip() for row in db.query(EstimateAssumption)
        .filter(EstimateAssumption.revision_id == rev.id)
        .order_by(EstimateAssumption.sort_order, EstimateAssumption.id).all()
        if row.text.strip()
    ]


def create_sow(db: Session, rev: EstimateRevision, user: User) -> SOW:
    if not sow_eligible(rev):
        raise ValueError("The MEP New Client SOW is available only for an approved Net New MEP estimate revision.")
    existing = latest_sow(db, rev.id)
    if existing:
        return existing
    tmpl = active_template(db)
    default_objective = (
        "Project will be executed with the objective of delivery and deployment of an automated solution "
        "to support Customer operations in Customers designated facilities and field operation."
    )
    sow = SOW(
        estimate_revision_id=rev.id,
        template_version_id=tmpl.id,
        sow_revision_no=1,
        status="DRAFT",
        sow_date=date.today(),
        agreement_type=AGREEMENT_TYPES[0],
        invoice_frequency="Weekly",
        project_objective=default_objective,
        created_by=user.id,
    )
    db.add(sow)
    db.flush()
    sites = max(int(rev.go_live_sites or 0), 0)
    for idx in range(sites):
        if rev.go_live_type == "On-Site All": support = "On-Site"
        elif rev.go_live_type == "On-Site Primary Remote Others": support = "On-Site" if idx == 0 else "Remote"
        else: support = "Remote"
        db.add(SOWHypercareLocation(sow_id=sow.id, support_type=support, sort_order=idx))
    record(db, event_type="SOW_CREATED", user_id=user.id, estimate_id=rev.estimate_id,
           revision_id=rev.id, field_name=f"SOW:{sow.id}", new_value="SOW Rev 1",
           reason=f"Pinned to {tmpl.label} v{tmpl.version_no}")
    db.commit()
    return sow


def copy_rejected_sow(db: Session, source: SOW, rev: EstimateRevision, user: User) -> SOW:
    if source.status != "REJECTED":
        raise ValueError("A new SOW revision can only be created from a Rejected SOW.")
    latest = latest_sow(db, source.estimate_revision_id)
    if latest and latest.id != source.id and latest.status in ("DRAFT", "FINALIZED", "PENDING_APPROVAL"):
        return latest
    data = {c.name: getattr(source, c.name) for c in SOW.__table__.columns if c.name not in {
        "id", "sow_revision_no", "status", "content_hash", "approved_text_snapshot",
        "created_by", "created_at", "updated_at", "finalized_by", "finalized_at", "submitted_by",
        "submitted_at", "approver_id", "approved_by", "approved_at", "rejected_by", "rejected_at", "rejection_reason"
    }}
    next_no = max(x.sow_revision_no for x in db.query(SOW).filter(SOW.estimate_revision_id == source.estimate_revision_id).all()) + 1
    data.update(sow_revision_no=next_no, status="DRAFT", created_by=user.id)
    dest = SOW(**data)
    db.add(dest); db.flush()
    for row in source.hypercare_locations:
        db.add(SOWHypercareLocation(sow_id=dest.id, description=row.description, country=row.country,
                                    support_type=row.support_type, allocated_hours=row.allocated_hours, sort_order=row.sort_order))
    for row in source.devices:
        db.add(SOWDevice(sow_id=dest.id, device_type=row.device_type, make_model=row.make_model,
                         os_version=row.os_version, sort_order=row.sort_order))
    record(db, event_type="SOW_REVISION_CREATED", user_id=user.id, estimate_id=rev.estimate_id,
           revision_id=rev.id, field_name=f"SOW:{dest.id}", old_value=f"SOW Rev {source.sow_revision_no}",
           new_value=f"SOW Rev {dest.sow_revision_no}", reason="Revision created from rejected SOW")
    db.commit()
    return dest


def validate_finalize(db: Session, sow: SOW, rev: EstimateRevision) -> list[str]:
    errors: list[str] = []
    if sow.agreement_type not in AGREEMENT_TYPES: errors.append("Select a valid Agreement Type.")
    if sow.invoice_frequency not in INVOICE_FREQUENCIES: errors.append("Select Weekly or Monthly invoice frequency.")
    if not sow.project_objective.strip(): errors.append("Project Objective is required.")
    if not sow.erp_version.strip(): errors.append("ERP Version is required for Appendix A.")
    if not sow.mep_product_version.strip(): errors.append("MEP Product Version is required for Appendix A.")
    if not sow.erp_deployment_model.strip(): errors.append("ERP Deployment Model is required for Appendix A.")
    if rev.erp == "Oracle JD Edwards E1":
        if not sow.erp_base_code_version.strip(): errors.append("ERP Base Code Version is required for a JDE EnterpriseOne SOW.")
        if not sow.erp_tools_release.strip(): errors.append("ERP Tools Release is required for a JDE EnterpriseOne SOW.")
    if rev.epp_install != "No":
        if not sow.epp_product_version.strip(): errors.append("EPP Product Version is required when EPP is included.")
        if not sow.print_methods.strip(): errors.append("Print Methods are required when EPP is included.")
    support_hours = go_live_support_hours(db, rev)
    allocated = sum(float(x.allocated_hours or 0) for x in sow.hypercare_locations)
    if abs(allocated - support_hours) > 0.01:
        errors.append(f"Hypercare allocations must equal the approved Go-Live Support hours ({support_hours:g}). Currently allocated: {allocated:g}.")
    if support_hours > 0:
        for idx, row in enumerate(sow.hypercare_locations, 1):
            if row.allocated_hours > 0 and not row.description.strip(): errors.append(f"Hypercare location {idx} needs a Location Description.")
            if row.allocated_hours > 0 and not row.country.strip(): errors.append(f"Hypercare location {idx} needs a Country.")
    return errors


def _replace_para_text(p: Paragraph, replacements: dict[str, str]) -> None:
    text = p.text
    if not any(k in text for k in replacements):
        return
    for k, v in replacements.items(): text = text.replace(k, v)
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]: r.text = ""
    else:
        p.add_run(text)


def _remove_conditional_blocks(doc: Document, include: dict[str, bool]) -> None:
    body = doc._element.body
    skip_key: str | None = None
    for el in list(body):
        if el.tag != qn("w:p"):
            if skip_key is not None:
                body.remove(el)
            continue
        text = "".join(t.text or "" for t in el.findall(".//" + qn("w:t"))).strip()
        if text.startswith("[[IF:") and text.endswith("]]" ):
            key = text[5:-2]
            skip_key = None if include.get(key, False) else key
            body.remove(el)
            continue
        if text.startswith("[[END:") and text.endswith("]]" ):
            key = text[6:-2]
            if skip_key == key: skip_key = None
            body.remove(el)
            continue
        if skip_key is not None:
            body.remove(el)


def _remove_paragraph(p: Paragraph) -> None:
    el = p._element
    if el.getparent() is not None: el.getparent().remove(el)


def _set_clone_text(el, text: str) -> None:
    ts = el.findall(".//" + qn("w:t"))
    if ts:
        ts[0].text = text
        for t in ts[1:]: t.text = ""


def _replace_list_marker(doc: Document, marker: str, values: list[str], lead_text: str | None = None) -> None:
    target = next((p for p in doc.paragraphs if marker in p.text), None)
    if not target:
        return
    if values:
        anchor = target._p
        for value in values:
            clone = deepcopy(target._p)
            _set_clone_text(clone, value)
            anchor.addprevious(clone)
    elif lead_text:
        lead = next((p for p in doc.paragraphs if p.text.strip() == lead_text), None)
        if lead: _remove_paragraph(lead)
    _remove_paragraph(target)


def _delete_table(table) -> None:
    el = table._element
    if el.getparent() is not None: el.getparent().remove(el)


def _delete_row(table, row) -> None:
    table._tbl.remove(row._tr)


def _fill_dynamic_tables(doc: Document, sow: SOW, rev: EstimateRevision, context: dict[str, str], support_hours: float) -> None:
    htable = next((t for t in doc.tables if t.rows and t.rows[0].cells[0].text.strip() == "Location Description"), None)
    if htable:
        while len(htable.rows) > 1: _delete_row(htable, htable.rows[-1])
        active = [x for x in sow.hypercare_locations if float(x.allocated_hours or 0) > 0 or x.description.strip()]
        if support_hours <= 0 or not active:
            _delete_table(htable)
        else:
            for row in active:
                cells = htable.add_row().cells
                cells[0].text = row.description
                cells[1].text = row.support_type
                cells[2].text = row.country
                cells[3].text = f"{float(row.allocated_hours or 0):g}"

    ctable = next((t for t in doc.tables if t.rows and "Approved Hourly Rate" in t.rows[0].cells[0].text), None)
    if ctable and len(ctable.rows) > 1:
        ctable.rows[1].cells[0].text = context["[[BILLING_RATE]]"]
        ctable.rows[1].cells[1].text = context["[[ESTIMATED_HOURS]]"]
        ctable.rows[1].cells[2].text = context["[[ESTIMATED_COST]]"]

    atable = next((t for t in doc.tables if t.rows and t.rows[0].cells[0].text.strip() == "Deployment Point"), None)
    if atable:
        token_values = {
            "[[ERP_SOLUTION]] [[ERP_VERSION]]": f"{rev.erp} {sow.erp_version}".strip(),
            "[[ERP_BASE_CODE]]": sow.erp_base_code_version,
            "[[ERP_TOOLS_RELEASE]]": sow.erp_tools_release,
            "[[ERP_OS]]": sow.erp_os_version,
            "[[ERP_DB]]": sow.erp_database_version,
            "[[MEP_VERSION]]": sow.mep_product_version,
            "[[EPP_VERSION]]": sow.epp_product_version if rev.epp_install != "No" else "",
            "[[PRINT_METHODS]]": sow.print_methods if rev.epp_install != "No" else "",
            "[[CI_DEPLOYMENT_MODEL]]": ci_deployment_model(rev),
            "[[ERP_DEPLOYMENT_MODEL]]": sow.erp_deployment_model,
        }
        for row in list(atable.rows):
            right = row.cells[1].text.strip() if len(row.cells) > 1 else ""
            for token, value in token_values.items():
                if token in right:
                    row.cells[1].text = value
                    if not value.strip(): _delete_row(atable, row)
                    break
        marker_index = next((i for i, row in enumerate(list(atable.rows)) if "[[DEVICE_ROWS]]" in row.cells[1].text), None)
        if marker_index is not None:
            for row in list(atable.rows)[marker_index:]:
                _delete_row(atable, row)
            for dev in sow.devices:
                if not dev.make_model.strip(): continue
                cells = atable.add_row().cells
                cells[0].text = dev.device_type
                cells[1].text = f"{dev.make_model} – {dev.os_version}" if dev.os_version.strip() else dev.make_model


def _all_doc_paragraphs(doc: Document):
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in doc.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def render_docx(db: Session, sow: SOW, rev: EstimateRevision) -> bytes:
    tmpl = db.get(SOWTemplateVersion, sow.template_version_id)
    if not tmpl: raise ValueError("The SOW template version no longer exists.")
    doc = Document(io.BytesIO(tmpl.content))
    support_hours = go_live_support_hours(db, rev)
    assumptions = estimate_assumptions(db, rev)
    baseline = selected_baseline_apps(db, rev)
    custom = selected_custom_apps(db, rev)
    labels = [f"Label {i}" for i in range(1, max(int(rev.label_count or 0), 0) + 1)]
    lines, _, _, _ = calculation(db, rev)
    load_hours = next((float(x.extended_hours or 0) for x in lines if x.key == "TEST_LOAD"), 0.0)
    include = {
        "MEP_CLOUD": rev.project_type in ("MEP Cloud", "Platform Move To Cloud"),
        "MEP_ON_PREM": rev.project_type in ("MEP On Prem", "Platform Move On Prem"),
        "EPP": rev.epp_install != "No" or rev.project_type in ("EPP Cloud", "EPP On Prem"),
        "REST_API": bool(sow.rest_api_required),
        "MIGRATE_TO_CLOUD": rev.project_type == "Platform Move To Cloud",
        "LIMITED_LOAD_TEST": load_hours > 0,
        "LOAD_BALANCER": bool(rev.high_availability),
        "ASSUMPTIONS": bool(assumptions),
        "JDE_REQUIREMENT": rev.erp == "Oracle JD Edwards E1",
        "LABELS": bool(labels),
    }
    _remove_conditional_blocks(doc, include)
    replacements = {
        "[[CUSTOMER_NAME]]": rev.customer or "",
        "[[ESTIMATE_NUMBER]]": rev.estimate.estimate_number,
        "[[SOW_DATE]]": sow.sow_date.strftime("%B %d, %Y"),
        "[[ENTITY]]": rev.entity or "Data Systems International, Inc. dba Cloud Inventory®",
        "[[AGREEMENT_TYPE]]": sow.agreement_type,
        "[[PROJECT_OBJECTIVE]]": sow.project_objective.strip(),
        "[[EPP_INSTALL_MODE]]": rev.epp_install,
        "[[BARCODE_PRINTER_COUNT]]": str(max(int(sow.barcode_printer_count or 0), 0)),
        "[[LABEL_COUNT]]": str(max(int(rev.label_count or 0), 0)),
        "[[CURRENCY]]": rev.currency,
        "[[BILLING_RATE]]": f"{rev.billing_rate:,.2f}",
        "[[ESTIMATED_HOURS]]": f"{rev.calculated_hours:g}",
        "[[ESTIMATED_COST]]": f"{rev.calculated_fees:,.2f}",
        "[[INVOICE_FREQUENCY]]": sow.invoice_frequency.lower(),
    }
    _replace_list_marker(doc, "[[BASELINE_APPLICATIONS]]", baseline, "Deploy and configure baseline")
    _replace_list_marker(doc, "[[CUSTOM_APPLICATIONS]]", custom, "Develop, deploy and configure personalized applications for Customer")
    _replace_list_marker(doc, "[[LABELS]]", labels)
    _replace_list_marker(doc, "[[ASSUMPTIONS]]", assumptions)
    for p in list(_all_doc_paragraphs(doc)):
        _replace_para_text(p, replacements)
    _fill_dynamic_tables(doc, sow, rev, replacements, support_hours)

    if sow.barcode_printer_count <= 0:
        for p in list(doc.paragraphs):
            if "Configuration and testing for up to 0 bar code label printers" in p.text:
                _remove_paragraph(p)

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        from docx.oxml import OxmlElement
        update = OxmlElement("w:updateFields"); settings.append(update)
    update.set(qn("w:val"), "true")
    out = io.BytesIO(); doc.save(out); return out.getvalue()


def canonical_text(docx_bytes: bytes) -> str:
    doc = Document(io.BytesIO(docx_bytes))
    chunks: list[str] = []
    def add_paragraphs(paragraphs):
        for p in paragraphs:
            text = p.text.strip()
            if text and not text.startswith("[[IF:") and not text.startswith("[[END:"):
                chunks.append(text)
    add_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            chunks.append(" | ".join(cell.text.strip() for cell in row.cells))
    for section in doc.sections:
        add_paragraphs(section.header.paragraphs)
        add_paragraphs(section.footer.paragraphs)
        for table in section.header.tables:
            for row in table.rows:
                chunks.append(" | ".join(cell.text.strip() for cell in row.cells))
        for table in section.footer.tables:
            for row in table.rows:
                chunks.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(chunks)


def content_hash_for(db: Session, sow: SOW, rev: EstimateRevision) -> tuple[str, str, bytes]:
    docx = render_docx(db, sow, rev)
    text = canonical_text(docx)
    return hashlib.sha256(text.encode("utf-8")).hexdigest(), text, docx


def verify_approved_content(db: Session, sow: SOW, rev: EstimateRevision) -> bytes:
    h, _, docx = content_hash_for(db, sow, rev)
    if sow.status == "APPROVED" and sow.content_hash and h != sow.content_hash:
        raise ValueError("The regenerated SOW no longer matches the approved wording. Historical download has been blocked for audit safety.")
    return docx


def render_pdf(db: Session, sow: SOW, rev: EstimateRevision) -> bytes:
    docx = verify_approved_content(db, sow, rev) if sow.status == "APPROVED" else render_docx(db, sow, rev)
    doc = Document(io.BytesIO(docx))
    out = io.BytesIO()
    pdf = SimpleDocTemplate(out, pagesize=letter, rightMargin=0.6*inch, leftMargin=0.6*inch, topMargin=0.65*inch, bottomMargin=0.65*inch)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="SOWTitle", parent=styles["Title"], alignment=TA_CENTER, fontSize=20, leading=24, spaceAfter=12))
    story = []
    if sow.status != "APPROVED":
        story += [RLParagraph(f"<b>{sow.status.replace('_',' ')}</b> — REVIEW COPY", styles["Normal"]), Spacer(1, 8)]
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            p = Paragraph(child, doc)
            text = p.text.strip()
            if not text: continue
            style_name = p.style.name if p.style else ""
            if text == "Statement Of Work": style = styles["SOWTitle"]
            elif style_name.startswith("Heading 1"): style = styles["Heading1"]
            elif style_name.startswith("Heading 2"): style = styles["Heading2"]
            elif style_name.startswith("Heading 3"): style = styles["Heading3"]
            else: style = styles["BodyText"]
            bullet = "•" if style_name == "List Paragraph" else None
            story.append(RLParagraph(html.escape(text), style, bulletText=bullet))
            story.append(Spacer(1, 3))
        elif child.tag == qn("w:tbl"):
            table = DocxTable(child, doc)
            data = [[RLParagraph(html.escape(cell.text.strip()), styles["BodyText"]) for cell in row.cells] for row in table.rows]
            if not data: continue
            rt = Table(data, repeatRows=1, hAlign="LEFT")
            rt.setStyle(TableStyle([
                ("GRID", (0,0), (-1,-1), 0.35, colors.grey),
                ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#e8eef1")),
                ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
                ("FONTSIZE", (0,0), (-1,-1), 8),
                ("VALIGN", (0,0), (-1,-1), "TOP"),
                ("LEFTPADDING", (0,0), (-1,-1), 4), ("RIGHTPADDING", (0,0), (-1,-1), 4),
            ]))
            story += [Spacer(1, 8), rt, Spacer(1, 8)]
    pdf.build(story)
    return out.getvalue()
