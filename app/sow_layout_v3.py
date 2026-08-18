from __future__ import annotations
import io
from datetime import datetime
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from sqlalchemy.orm import Session
from . import sow_layout_v2, sow_service
from .cip_models import PRODUCT_MEP
from .models import User
from .services.audit import record
from .sow_models import SOWTemplateVersion, SOW_TEMPLATE_MEP_NET_NEW

V3_FILENAME = 'MEP_Template_NewClient_2026_14_Controlled_v3.docx'
V3_REASON = ('SOW presentation correction: Heading 1/2/3 numbering standardized to X.0 / X.Y / X.Y.Z and '
             'the approver PDF review now reproduces the controlled cover, TOC, page starts, headers and footers.')
_PREVIOUS_SEED = sow_service.seed_initial_sow_template


def _build_v3_content(v2_content: bytes) -> bytes:
    doc = Document(io.BytesIO(v2_content))
    num_pr = doc.styles['Heading 1']._element.find('.//' + qn('w:numPr'))
    if num_pr is None or num_pr.find(qn('w:numId')) is None:
        raise RuntimeError('Controlled SOW Heading 1 is not linked to multilevel numbering.')
    num_id = num_pr.find(qn('w:numId')).get(qn('w:val'))
    root = doc.part.numbering_part.element
    num = next((x for x in root.findall(qn('w:num')) if x.get(qn('w:numId')) == num_id), None)
    if num is None: raise RuntimeError('Controlled SOW heading numbering instance is missing.')
    abstract_id = num.find(qn('w:abstractNumId')).get(qn('w:val'))
    abstract = next((x for x in root.findall(qn('w:abstractNum')) if x.get(qn('w:abstractNumId')) == abstract_id), None)
    if abstract is None: raise RuntimeError('Controlled SOW multilevel heading definition is missing.')
    formats = {0: '%1.0', 1: '%1.%2', 2: '%1.%2.%3'}
    for level in abstract.findall(qn('w:lvl')):
        ilvl = int(level.get(qn('w:ilvl')))
        if ilvl in formats:
            txt = level.find(qn('w:lvlText'))
            if txt is None:
                txt = OxmlElement('w:lvlText')
                level.append(txt)
            txt.set(qn('w:val'), formats[ilvl])
    for level, name in enumerate(('Heading 1', 'Heading 2', 'Heading 3')):
        style = doc.styles[name]
        style.paragraph_format.keep_with_next = True
        if level == 0: style.paragraph_format.page_break_before = True
        ppr = style._element.get_or_add_pPr(); npr = ppr.find(qn('w:numPr'))
        if npr is None: npr = OxmlElement('w:numPr'); ppr.append(npr)
        ilvl = npr.find(qn('w:ilvl'))
        if level:
            if ilvl is None: ilvl = OxmlElement('w:ilvl'); npr.insert(0, ilvl)
            ilvl.set(qn('w:val'), str(level))
        elif ilvl is not None: npr.remove(ilvl)
        nid = npr.find(qn('w:numId'))
        if nid is None: nid = OxmlElement('w:numId'); npr.append(nid)
        nid.set(qn('w:val'), num_id)
    # One source heading contains a direct numbering override; remove all such overrides so style numbering is authoritative.
    for p in doc.paragraphs:
        if p.style and p.style.name in ('Heading 1', 'Heading 2', 'Heading 3') and p._p.pPr is not None:
            direct = p._p.pPr.find(qn('w:numPr'))
            if direct is not None: p._p.pPr.remove(direct)
    appendix = next((p for p in doc.paragraphs if p.text.strip() == 'Appendix A'), None)
    if appendix:
        appendix.paragraph_format.page_break_before = True
        appendix.paragraph_format.keep_with_next = True
    settings = doc.settings._element; update = settings.find(qn('w:updateFields'))
    if update is None: update = OxmlElement('w:updateFields'); settings.append(update)
    update.set(qn('w:val'), 'true')
    for fld in doc._element.findall('.//' + qn('w:fldChar')):
        if fld.get(qn('w:fldCharType')) == 'begin': fld.set(qn('w:dirty'), 'true')
    out = io.BytesIO(); doc.save(out); return out.getvalue()


def seed_controlled_sow_template_v3(db: Session) -> None:
    _PREVIOUS_SEED(db)
    rows = db.query(SOWTemplateVersion).filter(SOWTemplateVersion.template_key == SOW_TEMPLATE_MEP_NET_NEW).order_by(SOWTemplateVersion.version_no).all()
    if any(row.version_no >= 3 for row in rows): return
    active = next((row for row in rows if row.status == 'ACTIVE'), None)
    if not active or not (active.version_no == 2 and active.filename == sow_layout_v2.V2_FILENAME and active.change_reason == sow_layout_v2.V2_REASON): return
    content = _build_v3_content(active.content); missing = sow_service.validate_template(content)
    if missing: raise RuntimeError(f"Controlled SOW template v3 is missing markers: {', '.join(missing)}")
    admin = db.query(User).filter(User.username_normalized == 'admin').first()
    if not admin: return
    now = datetime.utcnow(); active.status = 'RETIRED'; active.retired_at = now
    row = SOWTemplateVersion(template_key=SOW_TEMPLATE_MEP_NET_NEW, label='MEP New Client SOW', product_type=PRODUCT_MEP,
        customer_type='Net_New', version_no=3, status='ACTIVE', filename=V3_FILENAME, content=content,
        content_sha256=sow_service.sha256_bytes(content), change_reason=V3_REASON, created_by=admin.id,
        activated_by=admin.id, activated_at=now)
    db.add(row); db.flush()
    record(db, event_type='SOW_TEMPLATE_RETIRED', user_id=admin.id, field_name=f'SOW_TEMPLATE:{SOW_TEMPLATE_MEP_NET_NEW}:2',
           old_value=active.filename, new_value='RETIRED', reason='Superseded by controlled SOW presentation template v3.')
    record(db, event_type='SOW_TEMPLATE_ACTIVATED', user_id=admin.id, field_name=f'SOW_TEMPLATE:{SOW_TEMPLATE_MEP_NET_NEW}:3',
           new_value=row.filename, reason=row.change_reason)
    db.commit()

sow_service.seed_initial_sow_template = seed_controlled_sow_template_v3
from . import sow_pdf_v3  # noqa: E402,F401  # patches render_pdf after v3 seed is installed
