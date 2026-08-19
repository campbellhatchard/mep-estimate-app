from __future__ import annotations

import base64
import io
from datetime import date, datetime
from pathlib import Path

from docx import Document
from sqlalchemy import desc
from sqlalchemy.orm import Session

from ..assumptions import EstimateAssumption
from ..cip_models import CIPRevisionInput, CIPScopeItem, ConfigurationProduct, PRODUCT_CIP, PRODUCT_MEP
from ..models import ConfigItem, ConfigurationVersion, EstimateRevision, User
from ..services.audit import record
from ..services.cip_calculation import calculation as cip_calculation
from ..sow_models import SOW, SOWHypercareLocation, SOWTemplateVersion
from .. import sow_service

SOW_TEMPLATE_CIP_NET_NEW = "CIP_NET_NEW"
CIP_TEMPLATE_DIR = Path(__file__).parent / "sow_templates"
CIP_TEMPLATE_PART_GLOB = "cip_new_client_v1.b64.part*"
CIP_TEMPLATE_FILENAME = "CIP_Template_NewClient_2026_13_Controlled_v1.docx"
CIP_TEMPLATE_LABEL = "CIP New Client SOW"
CURRENT_VERSION_CATEGORY = "CIP SOW Setting"
CURRENT_VERSION_KEY = "CURRENT_CIP_VERSION"
CURRENT_VERSION_LABEL = "Current Version"

CIP_REQUIRED_TEMPLATE_MARKERS = {
    "[[CUSTOMER_NAME]]", "[[ESTIMATE_NUMBER]]", "[[SOW_DATE]]", "[[ENTITY]]",
    "[[AGREEMENT_TYPE]]", "[[PROJECT_OBJECTIVE]]", "[[EPP_INSTALL_MODE]]",
    "[[BARCODE_PRINTER_COUNT]]", "[[LABEL_COUNT]]", "[[BASELINE_APPLICATIONS]]",
    "[[CUSTOM_APPLICATIONS]]", "[[REPORTS]]", "[[INTEGRATIONS]]", "[[LABELS]]",
    "[[ASSUMPTIONS]]", "[[BILLING_RATE]]", "[[ESTIMATED_HOURS]]", "[[ESTIMATED_COST]]",
    "[[INVOICE_FREQUENCY]]", "[[ERP_SOLUTION]]", "[[ERP_VERSION]]", "[[ERP_BASE_CODE]]",
    "[[ERP_TOOLS_RELEASE]]", "[[ERP_OS]]", "[[ERP_DB]]", "[[CIP_VERSION]]",
    "[[EPP_VERSION]]", "[[PRINT_METHODS]]", "[[CI_DEPLOYMENT_MODEL]]",
    "[[ERP_DEPLOYMENT_MODEL]]", "[[DEVICE_ROWS]]", "[[HYPERCARE_ROWS]]", "[[CURRENCY]]",
    "[[IF:GATEWAY]]", "[[END:GATEWAY]]", "[[IF:EPP]]", "[[END:EPP]]",
    "[[IF:REST_API]]", "[[END:REST_API]]", "[[IF:LABELS]]", "[[END:LABELS]]",
    "[[IF:LIMITED_LOAD_TEST]]", "[[END:LIMITED_LOAD_TEST]]",
    "[[IF:ASSUMPTIONS]]", "[[END:ASSUMPTIONS]]",
    "[[IF:JDE_REQUIREMENT]]", "[[END:JDE_REQUIREMENT]]",
}

COVER_FOOTER = (
    "This statement of work estimate is the property and proprietary to Data Systems International, Inc. "
    "dba Cloud inventory® and contains trade secret and confidential information and is solely for Customer’s internal use.  "
    "Without the express written consent of Cloud Inventory ®, this estimate shall not be used, reproduced, copied, disclosed, "
    "transmitted in whole or in part.  Copyright © 2026 Data Systems International, Inc. dba Cloud Inventory®.  All rights reserved."
)


def _template_text(content: bytes) -> str:
    try:
        doc = Document(io.BytesIO(content))
    except Exception as exc:
        raise ValueError("The uploaded file is not a valid Word .docx document.") from exc
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    for section in doc.sections:
        chunks.extend(p.text for p in section.header.paragraphs)
        chunks.extend(p.text for p in section.footer.paragraphs)
    return "\n".join(chunks)


def validate_cip_template(content: bytes) -> list[str]:
    text = _template_text(content)
    return sorted(marker for marker in CIP_REQUIRED_TEMPLATE_MARKERS if marker not in text)


def validate_template_for_key(content: bytes, template_key: str) -> list[str]:
    if template_key == SOW_TEMPLATE_CIP_NET_NEW:
        return validate_cip_template(content)
    return sow_service.validate_template(content)


def _bundled_cip_template() -> bytes | None:
    parts = sorted(CIP_TEMPLATE_DIR.glob(CIP_TEMPLATE_PART_GLOB))
    if not parts:
        return None
    encoded = "".join(part.read_text() for part in parts)
    try:
        return base64.b64decode(encoded.strip())
    except Exception as exc:
        raise RuntimeError("Bundled CIP SOW template is not valid base64.") from exc


def ensure_cip_sow_data_settings(db: Session) -> None:
    versions = (
        db.query(ConfigurationVersion)
        .join(ConfigurationProduct, ConfigurationProduct.config_version_id == ConfigurationVersion.id)
        .filter(ConfigurationProduct.product_type == PRODUCT_CIP)
        .all()
    )
    changed = False
    for version in versions:
        exists = db.query(ConfigItem).filter(
            ConfigItem.config_version_id == version.id,
            ConfigItem.category == CURRENT_VERSION_CATEGORY,
            ConfigItem.key == CURRENT_VERSION_KEY,
        ).first()
        if exists:
            continue
        db.add(ConfigItem(
            config_version_id=version.id,
            category=CURRENT_VERSION_CATEGORY,
            key=CURRENT_VERSION_KEY,
            label=CURRENT_VERSION_LABEL,
            value_text=CURRENT_VERSION_LABEL,
            value_type="text",
            sort_order=0,
            active=True,
            description=(
                "CIP product version wording used for Net New CIP Statements of Work. "
                "Net New projects use the current production version."
            ),
        ))
        changed = True
    if changed:
        db.commit()


def seed_cip_sow_template(db: Session) -> None:
    ensure_cip_sow_data_settings(db)
    if db.query(SOWTemplateVersion).filter(
        SOWTemplateVersion.template_key == SOW_TEMPLATE_CIP_NET_NEW
    ).count():
        return
    content = _bundled_cip_template()
    if not content:
        return
    missing = validate_cip_template(content)
    if missing:
        raise RuntimeError("Bundled CIP SOW template is missing markers: " + ", ".join(missing))
    admin = db.query(User).filter(User.username_normalized == "admin").first()
    if not admin:
        return
    row = SOWTemplateVersion(
        template_key=SOW_TEMPLATE_CIP_NET_NEW,
        label=CIP_TEMPLATE_LABEL,
        product_type=PRODUCT_CIP,
        customer_type="Net_New",
        version_no=1,
        status="ACTIVE",
        filename=CIP_TEMPLATE_FILENAME,
        content=content,
        content_sha256=sow_service.sha256_bytes(content),
        change_reason="Initial controlled CIP New Client SOW template based on CIP_Template_NewClient_2026_13.",
        created_by=admin.id,
        activated_by=admin.id,
        activated_at=datetime.utcnow(),
    )
    db.add(row)
    db.flush()
    record(
        db,
        event_type="SOW_TEMPLATE_ACTIVATED",
        user_id=admin.id,
        field_name=f"SOW_TEMPLATE:{SOW_TEMPLATE_CIP_NET_NEW}:1",
        new_value=row.filename,
        reason=row.change_reason,
    )
    db.commit()


def _active_cip_template(db: Session) -> SOWTemplateVersion:
    row = (
        db.query(SOWTemplateVersion)
        .filter(
            SOWTemplateVersion.template_key == SOW_TEMPLATE_CIP_NET_NEW,
            SOWTemplateVersion.status == "ACTIVE",
        )
        .order_by(desc(SOWTemplateVersion.version_no))
        .first()
    )
    if not row:
        raise ValueError("No active CIP New Client SOW template is available.")
    return row


def cip_sow_eligible(rev: EstimateRevision) -> bool:
    return (
        rev.status in ("APPROVED", "FINAL", "SUPERSEDED")
        and bool((rev.engine_version or "").upper().startswith("CIP-"))
        and rev.customer_type == "Net_New"
        and rev.project_type == "CIP Install"
    )


def current_cip_version(db: Session, rev: EstimateRevision) -> str:
    row = db.query(ConfigItem).filter(
        ConfigItem.config_version_id == rev.config_version_id,
        ConfigItem.category == CURRENT_VERSION_CATEGORY,
        ConfigItem.key == CURRENT_VERSION_KEY,
        ConfigItem.active.is_(True),
    ).first()
    return (row.value_text or row.label).strip() if row else CURRENT_VERSION_LABEL


def create_cip_sow(db: Session, rev: EstimateRevision, user: User) -> SOW:
    if not cip_sow_eligible(rev):
        raise ValueError("The CIP New Client SOW is available only for an approved Net New CIP Install estimate revision.")
    existing = sow_service.latest_sow(db, rev.id)
    if existing:
        return existing
    inp = db.get(CIPRevisionInput, rev.id)
    if not inp:
        raise ValueError("CIP estimate inputs are not available.")
    tmpl = _active_cip_template(db)
    sow = SOW(
        estimate_revision_id=rev.id,
        template_version_id=tmpl.id,
        sow_revision_no=1,
        status="DRAFT",
        sow_date=date.today(),
        agreement_type=sow_service.AGREEMENT_TYPES[0],
        invoice_frequency="Weekly",
        project_objective=(
            "Project will be executed with the objective of delivery and deployment of an automated solution "
            "to support Customer operations in Customers designated facilities and field operation."
        ),
        rest_api_required=bool(inp.rest_required),
        mep_product_version=current_cip_version(db, rev),
        created_by=user.id,
    )
    db.add(sow)
    db.flush()
    sites = max(int(inp.go_live_sites or 0), 0)
    for idx in range(sites):
        if inp.go_live_type == "On-Site All":
            support = "On-Site"
        elif inp.go_live_type == "On-Site Primary Remote Others":
            support = "On-Site" if idx == 0 else "Remote"
        else:
            support = "Remote"
        db.add(SOWHypercareLocation(sow_id=sow.id, support_type=support, sort_order=idx))
    record(
        db,
        event_type="SOW_CREATED",
        user_id=user.id,
        estimate_id=rev.estimate_id,
        revision_id=rev.id,
        field_name=f"SOW:{sow.id}",
        new_value="SOW Rev 1",
        reason=f"Pinned to {tmpl.label} v{tmpl.version_no}; CIP version {sow.mep_product_version}",
    )
    db.commit()
    return sow


def cip_go_live_support_hours(db: Session, rev: EstimateRevision) -> float:
    lines, _, _, _ = cip_calculation(db, rev)
    for line in lines:
        if getattr(line, "key", "") == "GOLIVE_SUPPORT":
            return float(getattr(line, "investment_hours", 0) or 0)
    return 0.0


def _scope_rows(db: Session, rev: EstimateRevision, category: str) -> list[CIPScopeItem]:
    return (
        db.query(CIPScopeItem)
        .filter(CIPScopeItem.revision_id == rev.id, CIPScopeItem.category == category)
        .order_by(CIPScopeItem.sort_order, CIPScopeItem.id)
        .all()
    )


def _selected(rows: list[CIPScopeItem]) -> list[CIPScopeItem]:
    return [row for row in rows if row.config_type != "No Config"]


def cip_scope_lists(db: Session, rev: EstimateRevision) -> dict[str, list[str]]:
    desktop = _selected(_scope_rows(db, rev, "DESKTOP"))
    mobile = _selected(_scope_rows(db, rev, "MOBILE"))
    baseline = [row.label for row in desktop + mobile if row.config_type == "Baseline"]
    modified = [f"{row.label} — Mod Required" for row in desktop + mobile if row.config_type == "Mod Required"]

    custom = []
    for category in ("CUSTOM_DESKTOP", "CUSTOM_MOBILE"):
        for row in _selected(_scope_rows(db, rev, category)):
            custom.append((row.description or row.label).strip())
    custom = modified + [x for x in custom if x]

    reports = [
        (row.description or row.label).strip()
        for row in _selected(_scope_rows(db, rev, "REPORT"))
        if (row.description or row.label).strip()
    ]

    integrations = [row.label for row in _selected(_scope_rows(db, rev, "INTEGRATION"))]
    for category in ("CUSTOM_BOOMI", "REST"):
        for row in _selected(_scope_rows(db, rev, category)):
            value = (row.description or row.label).strip()
            if value:
                integrations.append(value)

    labels = [
        (row.description or row.label).strip()
        for row in _selected(_scope_rows(db, rev, "LABEL"))
        if (row.description or row.label).strip()
    ]
    return {
        "baseline": baseline,
        "custom": custom,
        "reports": reports,
        "integrations": integrations,
        "labels": labels,
    }


def _assumptions(db: Session, rev: EstimateRevision) -> list[str]:
    return [
        row.text.strip()
        for row in db.query(EstimateAssumption)
        .filter(EstimateAssumption.revision_id == rev.id)
        .order_by(EstimateAssumption.sort_order, EstimateAssumption.id)
        .all()
        if row.text.strip()
    ]
