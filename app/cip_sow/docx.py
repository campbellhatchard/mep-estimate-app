from __future__ import annotations

import hashlib
import io

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from sqlalchemy.orm import Session

from ..cip_models import CIPRevisionInput
from ..models import EstimateRevision
from ..services.cip_calculation import calculation as cip_calculation
from ..sow_models import SOW, SOWTemplateVersion
from .. import sow_service
from .core import SOW_TEMPLATE_CIP_NET_NEW, cip_go_live_support_hours, cip_scope_lists, _assumptions


def _delete_table(table) -> None:
    el = table._element
    if el.getparent() is not None:
        el.getparent().remove(el)


def _delete_row(table, row) -> None:
    table._tbl.remove(row._tr)


def _fill_cip_tables(
    doc: Document,
    sow: SOW,
    rev: EstimateRevision,
    inp: CIPRevisionInput,
    replacements: dict[str, str],
    support_hours: float,
) -> None:
    htable = next(
        (t for t in doc.tables if t.rows and t.rows[0].cells[0].text.strip() == "Location Description"),
        None,
    )
    if htable:
        while len(htable.rows) > 1:
            _delete_row(htable, htable.rows[-1])
        active = [
            x for x in sow.hypercare_locations
            if float(x.allocated_hours or 0) > 0 or x.description.strip()
        ]
        if support_hours <= 0 or not active:
            _delete_table(htable)
        else:
            for row in active:
                cells = htable.add_row().cells
                cells[0].text = row.description
                cells[1].text = row.support_type
                cells[2].text = row.country
                cells[3].text = f"{float(row.allocated_hours or 0):g}"

    ctable = next(
        (t for t in doc.tables if t.rows and "Approved Hourly Rate" in t.rows[0].cells[0].text),
        None,
    )
    if ctable and len(ctable.rows) > 1:
        ctable.rows[1].cells[0].text = replacements["[[BILLING_RATE]]"]
        ctable.rows[1].cells[1].text = replacements["[[ESTIMATED_HOURS]]"]
        ctable.rows[1].cells[2].text = replacements["[[ESTIMATED_COST]]"]

    atable = next(
        (t for t in doc.tables if t.rows and t.rows[0].cells[0].text.strip() == "Deployment Point"),
        None,
    )
    if not atable:
        return

    standalone = inp.deployed_over == "Standalone"
    jde = inp.deployed_over == "JD Edwards"
    epp = inp.epp_install != "No"

    for row in list(atable.rows):
        left = row.cells[0].text.strip() if row.cells else ""
        right = row.cells[1].text.strip() if len(row.cells) > 1 else ""

        if standalone and any(label in left for label in (
            "ERP Base Code Version", "ERP Tools Release Version", "ERP Operating System Version",
            "ERP Database Type / Version", "Deployment Model – ERP Solution",
        )):
            _delete_row(atable, row)
            continue
        if not jde and any(label in left for label in ("ERP Base Code Version", "ERP Tools Release Version")):
            _delete_row(atable, row)
            continue
        if not epp and any(label in left for label in ("Label Printing Solution", "Print Method(s)")):
            _delete_row(atable, row)
            continue

        token_values = {
            "[[ERP_SOLUTION]] [[ERP_VERSION]]": (
                inp.deployed_over if standalone else f"{inp.deployed_over} {sow.erp_version}".strip()
            ),
            "[[ERP_BASE_CODE]]": sow.erp_base_code_version,
            "[[ERP_TOOLS_RELEASE]]": sow.erp_tools_release,
            "[[ERP_OS]]": sow.erp_os_version,
            "[[ERP_DB]]": sow.erp_database_version,
            "[[CIP_VERSION]]": sow.mep_product_version,
            "[[EPP_VERSION]]": sow.epp_product_version if epp else "",
            "[[PRINT_METHODS]]": sow.print_methods if epp else "",
            "[[CI_DEPLOYMENT_MODEL]]": "Cloud Inventory® Managed / Public Cloud",
            "[[ERP_DEPLOYMENT_MODEL]]": sow.erp_deployment_model if not standalone else "",
        }
        for token, value in token_values.items():
            if token in right:
                row.cells[1].text = value
                if not value.strip():
                    _delete_row(atable, row)
                break

    # python-docx creates fresh _Row wrappers each time table.rows is accessed, so
    # determine the marker index from one stable list rather than comparing wrappers.
    rows = list(atable.rows)
    marker_index = next(
        (
            idx for idx, row in enumerate(rows)
            if len(row.cells) > 1 and "[[DEVICE_ROWS]]" in row.cells[1].text
        ),
        None,
    )
    if marker_index is not None:
        for row in rows[marker_index:]:
            _delete_row(atable, row)
        for dev in sow.devices:
            if not dev.make_model.strip():
                continue
            cells = atable.add_row().cells
            cells[0].text = dev.device_type
            cells[1].text = (
                f"{dev.make_model} – {dev.os_version}" if dev.os_version.strip() else dev.make_model
            )


def render_cip_docx(db: Session, sow: SOW, rev: EstimateRevision) -> bytes:
    template = db.get(SOWTemplateVersion, sow.template_version_id)
    if not template or template.template_key != SOW_TEMPLATE_CIP_NET_NEW:
        raise ValueError("The CIP SOW template version no longer exists.")
    inp = db.get(CIPRevisionInput, rev.id)
    if not inp:
        raise ValueError("CIP estimate inputs are not available.")

    doc = Document(io.BytesIO(template.content))
    support_hours = cip_go_live_support_hours(db, rev)
    assumptions = _assumptions(db, rev)
    scope = cip_scope_lists(db, rev)
    calc_lines, _, _, _ = cip_calculation(db, rev)
    limited_load_hours = next(
        (
            float(getattr(x, "investment_hours", 0) or 0)
            for x in calc_lines
            if getattr(x, "key", "") == "TEST_LIMITED_LOAD"
        ),
        0.0,
    )

    include = {
        "GATEWAY": bool(inp.gateway),
        "EPP": inp.epp_install != "No",
        "REST_API": bool(inp.rest_required and inp.rest_interface_count > 0),
        "LABELS": bool(scope["labels"]),
        "LIMITED_LOAD_TEST": limited_load_hours > 0,
        "ASSUMPTIONS": bool(assumptions),
        "JDE_REQUIREMENT": inp.deployed_over == "JD Edwards",
    }
    sow_service._remove_conditional_blocks(doc, include)

    replacements = {
        "[[CUSTOMER_NAME]]": rev.customer or "",
        "[[ESTIMATE_NUMBER]]": rev.estimate.estimate_number,
        "[[SOW_DATE]]": sow.sow_date.strftime("%B %d, %Y"),
        "[[ENTITY]]": rev.entity or "Data Systems International, Inc. dba Cloud Inventory®",
        "[[AGREEMENT_TYPE]]": sow.agreement_type,
        "[[PROJECT_OBJECTIVE]]": sow.project_objective.strip(),
        "[[EPP_INSTALL_MODE]]": inp.epp_install,
        "[[BARCODE_PRINTER_COUNT]]": str(max(int(sow.barcode_printer_count or 0), 0)),
        "[[LABEL_COUNT]]": str(max(int(inp.label_count or 0), 0)),
        "[[CURRENCY]]": rev.currency,
        "[[BILLING_RATE]]": f"{rev.billing_rate:,.2f}",
        "[[ESTIMATED_HOURS]]": f"{rev.calculated_hours:g}",
        "[[ESTIMATED_COST]]": f"{rev.calculated_fees:,.2f}",
        "[[INVOICE_FREQUENCY]]": sow.invoice_frequency.lower(),
    }

    sow_service._replace_list_marker(doc, "[[BASELINE_APPLICATIONS]]", scope["baseline"], "Deploy and configure baseline")
    sow_service._replace_list_marker(
        doc, "[[CUSTOM_APPLICATIONS]]", scope["custom"],
        "Develop, deploy and configure personalized applications for Customer",
    )
    sow_service._replace_list_marker(
        doc, "[[REPORTS]]", scope["reports"], "Develop, deploy and configure reports for Customer",
    )
    sow_service._replace_list_marker(doc, "[[INTEGRATIONS]]", scope["integrations"])
    sow_service._replace_list_marker(doc, "[[LABELS]]", scope["labels"])
    sow_service._replace_list_marker(doc, "[[ASSUMPTIONS]]", assumptions)

    for paragraph in list(sow_service._all_doc_paragraphs(doc)):
        sow_service._replace_para_text(paragraph, replacements)

    _fill_cip_tables(doc, sow, rev, inp, replacements, support_hours)

    if sow.barcode_printer_count <= 0:
        for p in list(doc.paragraphs):
            if "Configuration and testing for up to 0 bar code label printers" in p.text:
                sow_service._remove_paragraph(p)

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")
    for fld in doc._element.findall(".//" + qn("w:fldChar")):
        if fld.get(qn("w:fldCharType")) == "begin":
            fld.set(qn("w:dirty"), "true")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def cip_content_hash_for(db: Session, sow: SOW, rev: EstimateRevision) -> tuple[str, str, bytes]:
    docx = render_cip_docx(db, sow, rev)
    text = sow_service.canonical_text(docx)
    return hashlib.sha256(text.encode("utf-8")).hexdigest(), text, docx


def verify_cip_approved_content(db: Session, sow: SOW, rev: EstimateRevision) -> bytes:
    digest, _, docx = cip_content_hash_for(db, sow, rev)
    if sow.status == "APPROVED" and sow.content_hash and digest != sow.content_hash:
        raise ValueError(
            "The regenerated SOW no longer matches the approved wording. "
            "Historical download has been blocked for audit safety."
        )
    return docx


def validate_cip_finalize(db: Session, sow: SOW, rev: EstimateRevision) -> list[str]:
    inp = db.get(CIPRevisionInput, rev.id)
    if not inp:
        return ["CIP estimate inputs are not available."]
    errors: list[str] = []
    if sow.agreement_type not in sow_service.AGREEMENT_TYPES:
        errors.append("Select a valid Agreement Type.")
    if sow.invoice_frequency not in sow_service.INVOICE_FREQUENCIES:
        errors.append("Select Weekly or Monthly invoice frequency.")
    if not sow.project_objective.strip():
        errors.append("Project Objective is required.")
    if not sow.mep_product_version.strip():
        errors.append("CIP Product Version is required for Appendix A.")
    if sow.barcode_printer_count < 0:
        errors.append("Barcode Printer Count cannot be negative.")

    if inp.deployed_over != "Standalone":
        if not sow.erp_version.strip():
            errors.append("ERP/System Version is required for Appendix A.")
        if not sow.erp_deployment_model.strip():
            errors.append("ERP Deployment Model is required for Appendix A.")
    if inp.deployed_over == "JD Edwards":
        if not sow.erp_base_code_version.strip():
            errors.append("ERP Base Code Version is required for a JD Edwards SOW.")
        if not sow.erp_tools_release.strip():
            errors.append("ERP Tools Release is required for a JD Edwards SOW.")

    if inp.epp_install != "No":
        if not sow.epp_product_version.strip():
            errors.append("EPP Product Version is required when EPP is included.")
        if not sow.print_methods.strip():
            errors.append("Print Methods are required when EPP is included.")

    support_hours = cip_go_live_support_hours(db, rev)
    allocated = sum(float(x.allocated_hours or 0) for x in sow.hypercare_locations)
    if abs(allocated - support_hours) > 0.01:
        errors.append(
            f"Hypercare allocations must equal the approved Go-Live Support hours "
            f"({support_hours:g}). Currently allocated: {allocated:g}."
        )
    if support_hours > 0:
        for idx, row in enumerate(sow.hypercare_locations, 1):
            if row.allocated_hours > 0 and not row.description.strip():
                errors.append(f"Hypercare location {idx} needs a Location Description.")
            if row.allocated_hours > 0 and not row.country.strip():
                errors.append(f"Hypercare location {idx} needs a Country.")
    return errors
