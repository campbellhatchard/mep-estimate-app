from __future__ import annotations

import base64
import io
import json
import zlib
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from sqlalchemy import desc
from sqlalchemy.orm import Session

from ..sow_models import SOWTemplateVersion

SPEC_DIR = Path(__file__).parent / "template_spec"
SPEC_GLOB = "cip_new_client_spec.b64.part*"


def _spec() -> list:
    parts = sorted(SPEC_DIR.glob(SPEC_GLOB))
    if not parts:
        raise RuntimeError("Bundled CIP SOW template specification is missing.")
    encoded = "".join(part.read_text() for part in parts)
    try:
        raw = zlib.decompress(base64.b64decode(encoded.strip()))
        return json.loads(raw.decode("utf-8"))
    except Exception as exc:
        raise RuntimeError("Bundled CIP SOW template specification is invalid.") from exc


def _text_of(el) -> str:
    return "".join(node.text or "" for node in el.findall(".//" + qn("w:t")))


def _find_toc(body):
    for child in body:
        instr = " ".join((n.text or "") for n in child.findall(".//" + qn("w:instrText")))
        if "TOC" in instr:
            return deepcopy(child)
    raise RuntimeError("Controlled MEP template is missing its Word TOC field.")


def _table_key(table: DocxTable) -> str:
    first = table.rows[0].cells[0].text.strip() if table.rows and table.rows[0].cells else ""
    if first == "Location Description":
        return "hypercare"
    if first == "Approved Hourly Rate":
        return "cost"
    if first == "Deployment Point":
        return "appendix"
    if first.startswith("By execution, signer certifies"):
        return "signature"
    return ""


def _prototypes(doc: Document) -> dict[str, object]:
    result = {}
    for table in doc.tables:
        key = _table_key(table)
        if key:
            result[key] = deepcopy(table._tbl)
    missing = {"hypercare", "cost", "signature", "appendix"} - set(result)
    if missing:
        raise RuntimeError("Controlled MEP template is missing reusable SOW table layout(s): " + ", ".join(sorted(missing)))
    return result


def _clear_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _append_paragraph(doc: Document, style_name: str, text: str, page_break: bool) -> None:
    try:
        p = doc.add_paragraph(style=style_name or None)
    except KeyError:
        p = doc.add_paragraph()
    if text:
        p.add_run(text)
    if page_break:
        p.add_run().add_break(WD_BREAK.PAGE)


def _clone_table_for_rows(doc: Document, proto, rows: list[list[str]]) -> None:
    tbl = deepcopy(proto)
    existing = tbl.findall(qn("w:tr"))
    if not existing:
        return
    while len(existing) < len(rows):
        tbl.append(deepcopy(existing[-1]))
        existing = tbl.findall(qn("w:tr"))
    while len(existing) > len(rows):
        tbl.remove(existing[-1])
        existing = tbl.findall(qn("w:tr"))
    doc._element.body.insert(len(doc._element.body) - 1, tbl)
    table = DocxTable(tbl, doc)
    for r_idx, row_values in enumerate(rows):
        for c_idx, value in enumerate(row_values):
            if c_idx < len(table.rows[r_idx].cells):
                table.rows[r_idx].cells[c_idx].text = value


def _prototype_for_spec(rows: list[list[str]], protos: dict[str, object]):
    first = rows[0][0].strip() if rows and rows[0] else ""
    if first == "Location Description":
        return protos["hypercare"]
    if first == "Approved Hourly Rate":
        return protos["cost"]
    if first == "Deployment Point":
        return protos["appendix"]
    if first.startswith("By execution, signer certifies"):
        return protos["signature"]
    raise RuntimeError(f"Unexpected CIP SOW table in specification: {first[:80]}")


def build_cip_template(db: Session) -> bytes:
    """Build CIP New Client template from the approved CIP source spec using the accepted MEP house layout."""
    mep = (
        db.query(SOWTemplateVersion)
        .filter(SOWTemplateVersion.template_key == "MEP_NET_NEW", SOWTemplateVersion.status == "ACTIVE")
        .order_by(desc(SOWTemplateVersion.version_no))
        .first()
    )
    if not mep:
        raise RuntimeError("An active controlled MEP SOW template is required to seed the CIP template.")
    doc = Document(io.BytesIO(mep.content))
    toc = _find_toc(doc._element.body)
    protos = _prototypes(doc)
    _clear_body(doc)

    toc_inserted = False

    def _insert_toc() -> None:
        nonlocal toc_inserted
        if not toc_inserted:
            doc._element.body.insert(len(doc._element.body) - 1, deepcopy(toc))
            toc_inserted = True

    for item in _spec():
        kind = item[0]
        if kind == "p":
            _, style_name, text, page_break, special = item
            if special == "TOC":
                _insert_toc()
            else:
                # The bundled spec does not carry an explicit TOC marker item (every
                # entry's `special` field is empty), so the field is never placed by
                # the branch above. Insert it at the position a real Word
                # "Insert > Table of Contents" page occupies: immediately before the
                # first numbered heading, i.e. right before Section 1.0. Without
                # this, the built CIP template contains no TOC field at all, and the
                # controlled-Word TOC/page-number reconciliation in
                # app/sow_word_control.py fails every CIP SOW with a 409.
                if style_name == "Heading 1":
                    _insert_toc()
                _append_paragraph(doc, style_name, text, bool(page_break))
        elif kind == "t":
            rows = item[1]
            _clone_table_for_rows(doc, _prototype_for_spec(rows, protos), rows)
        elif kind == "sect":
            continue

    # Defensive: guarantee a TOC field is always present even if some future
    # revision of the bundled spec has no Heading 1 at all.
    _insert_toc()

    # Preserve the accepted template behavior: Word refreshes TOC and page fields on open.
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        from docx.oxml import OxmlElement
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")
    for fld in doc._element.findall(".//" + qn("w:fldChar")):
        if fld.get(qn("w:fldCharType")) == "begin":
            fld.set(qn("w:dirty"), "true")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
