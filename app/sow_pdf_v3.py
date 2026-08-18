from __future__ import annotations
import html, io, zipfile
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table as DocxTable
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen.canvas import Canvas
from reportlab.platypus import BaseDocTemplate, Frame, PageBreak, PageTemplate, Paragraph as RP, Spacer, Table, TableStyle
from reportlab.platypus.tableofcontents import TableOfContents
from . import sow_service
from .sow_models import SOWTemplateVersion

COVER_FOOTER = ('This statement of work estimate is the property and proprietary to Data Systems International, Inc. '
    'dba Cloud inventory® and contains trade secret and confidential information and is solely for Customer’s internal use.  '
    'Without the express written consent of Cloud Inventory ®, this estimate shall not be used, reproduced, copied, disclosed, '
    'transmitted in whole or in part.  Copyright © 2026 Data Systems International, Inc. dba Cloud Inventory®.  All rights reserved.')
_PREVIOUS_RENDER_PDF = sow_service.render_pdf


def _logo(docx_bytes: bytes):
    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
            names = [n for n in z.namelist() if n.startswith('word/media/') and n.lower().endswith(('.png','.jpg','.jpeg'))]
            return z.read(names[0]) if names else None
    except Exception:
        return None


class _Doc(BaseDocTemplate):
    def __init__(self, *a, **k): super().__init__(*a, **k); self._seq = 0
    def beforeDocument(self): self._seq = 0; return super().beforeDocument()
    def afterFlowable(self, f):
        if isinstance(f, RP) and hasattr(f, '_toc_level'):
            self._seq += 1; key = f'sow_h_{self._seq}'; self.canv.bookmarkPage(key)
            self.notify('TOCEntry', (f._toc_level, f.getPlainText(), self.page, key))


def _canvas(customer: str, estimate: str, logo_bytes: bytes | None):
    class _Canvas(Canvas):
        def __init__(self,*a,**k):
            Canvas.__init__(self,*a,**k); self._states=[]; self._logo=ImageReader(io.BytesIO(logo_bytes)) if logo_bytes else None
        def showPage(self): self._states.append(dict(self.__dict__)); self._startPage()
        def save(self):
            total=len(self._states)
            for state in self._states:
                self.__dict__.update(state); self._decorate(total); Canvas.showPage(self)
            Canvas.save(self)
        def _decorate(self,total):
            w,h=letter; left=.72*inch; right=w-.72*inch
            if self._logo:
                iw,ih=self._logo.getSize(); width=120; height=width*ih/iw
                self.drawImage(self._logo,left,h-.28*inch-height,width=width,height=height,mask='auto',preserveAspectRatio=True)
            if self._pageNumber>1:
                self.setFont('Helvetica-Bold',9); self.drawRightString(right,h-.82*inch,customer)
            self.setLineWidth(1); self.line(left,h-1.03*inch,right,h-1.03*inch)
            if self._pageNumber==1:
                st=ParagraphStyle('coverFooter',fontName='Helvetica',fontSize=6.3,leading=7.2)
                p=RP(html.escape(COVER_FOOTER),st); p.wrap(right-left,.6*inch); p.drawOn(self,left,.28*inch)
            else:
                self.setFont('Helvetica-Bold',7.5); self.drawString(left,.42*inch,f'Estimate Number: {estimate}')
                self.drawCentredString(w/2,.42*inch,f'Page {self._pageNumber} of {total}')
    return _Canvas


def _review_pdf(docx_bytes: bytes, customer: str, estimate: str, sow_date: str) -> bytes:
    doc=Document(io.BytesIO(docx_bytes)); out=io.BytesIO(); frame=Frame(.72*inch,.72*inch,letter[0]-1.44*inch,letter[1]-1.78*inch,id='body',topPadding=0,bottomPadding=0)
    pdf=_Doc(out,pagesize=letter); pdf.addPageTemplates([PageTemplate(id='sow',frames=[frame])]); s=getSampleStyleSheet()
    def add(name,parent,**kw): s.add(ParagraphStyle(name,parent=s[parent],**kw))
    add('CoverTitle','Title',fontName='Helvetica-Bold',fontSize=20,leading=24,alignment=TA_CENTER,spaceAfter=24)
    add('CoverCustomer','Normal',fontName='Helvetica-Bold',fontSize=16,leading=20,alignment=TA_CENTER,spaceAfter=60)
    add('CoverSubtitle','Normal',fontName='Helvetica-Bold',fontSize=12.5,leading=16,alignment=TA_CENTER,spaceAfter=22)
    add('CoverMeta','Normal',fontName='Helvetica-Bold',fontSize=11,leading=14,alignment=TA_CENTER,spaceAfter=28)
    add('TOCTitle','Heading1',fontName='Helvetica-Bold',fontSize=16,leading=20,alignment=TA_CENTER,spaceAfter=18)
    add('H1S','Heading1',fontName='Helvetica-Bold',fontSize=14,leading=17,spaceAfter=10,keepWithNext=True)
    add('H2S','Heading2',fontName='Helvetica-Bold',fontSize=11.5,leading=14,spaceBefore=7,spaceAfter=5,keepWithNext=True)
    add('H3S','Heading3',fontName='Helvetica-Bold',fontSize=10.5,leading=13,spaceBefore=5,spaceAfter=4,keepWithNext=True)
    add('BodyS','BodyText',fontName='Helvetica',fontSize=8.7,leading=11.3,spaceAfter=4)
    add('BulletS','BodyS',leftIndent=18,firstLineIndent=-8,bulletIndent=6,spaceAfter=2)
    add('AppendixS','Heading1',fontName='Helvetica-Bold',fontSize=16,leading=20,alignment=TA_CENTER,spaceAfter=10,keepWithNext=True)
    story=[Spacer(1,1.28*inch),RP('Statement Of Work',s['CoverTitle']),RP(html.escape(customer),s['CoverCustomer']),
        RP('Delivery Estimate for Service to Deploy and Configure Mobile Enterprise<br/>Platform Solution',s['CoverSubtitle']),
        RP(f'Estimate Number: {html.escape(estimate)}',s['CoverMeta']),RP(f'Date: {html.escape(sow_date)}',s['CoverMeta']),PageBreak(),RP('Table of Contents',s['TOCTitle'])]
    toc=TableOfContents(); toc.levelStyles=[ParagraphStyle('T1',fontName='Helvetica-Bold',fontSize=9.2,leading=12,leftIndent=0,spaceBefore=4),ParagraphStyle('T2',fontName='Helvetica',fontSize=8.5,leading=11,leftIndent=14),ParagraphStyle('T3',fontName='Helvetica',fontSize=8.2,leading=10,leftIndent=28)]
    story += [toc,PageBreak()]
    children=list(doc._element.body); toc_i=None
    for i,ch in enumerate(children):
        if 'TOC' in ' '.join((n.text or '') for n in ch.findall('.//'+qn('w:instrText'))): toc_i=i; break
    counters=[0,0,0]
    for ch in children[(toc_i+1 if toc_i is not None else 0):]:
        if ch.tag==qn('w:p'):
            p=Paragraph(ch,doc); text=p.text.strip(); style=p.style.name if p.style else ''
            if not text: continue
            if text=='Appendix A':
                story.append(PageBreak()); f=RP('Appendix A',s['AppendixS']); f._toc_level=0; story.append(f); continue
            if style=='Heading 1':
                counters=[counters[0]+1,0,0]; story.append(PageBreak()); f=RP(f'{counters[0]}.0 {html.escape(text)}',s['H1S']); f._toc_level=0; story.append(f)
            elif style=='Heading 2':
                counters[1]+=1; counters[2]=0; f=RP(f'{counters[0]}.{counters[1]} {html.escape(text)}',s['H2S']); f._toc_level=1; story.append(f)
            elif style=='Heading 3':
                counters[2]+=1; f=RP(f'{counters[0]}.{counters[1]}.{counters[2]} {html.escape(text)}',s['H3S']); f._toc_level=2; story.append(f)
            elif not text.startswith('[[IF:') and not text.startswith('[[END:'):
                story.append(RP(html.escape(text),s['BulletS'] if style=='List Paragraph' else s['BodyS'],bulletText='•' if style=='List Paragraph' else None))
        elif ch.tag==qn('w:tbl'):
            src=DocxTable(ch,doc); data=[[RP(html.escape(c.text.strip()),s['BodyS']) for c in r.cells] for r in src.rows]
            if data:
                t=Table(data,repeatRows=1,hAlign='LEFT'); t.setStyle(TableStyle([('GRID',(0,0),(-1,-1),.45,colors.HexColor('#7d8d94')),('BACKGROUND',(0,0),(-1,0),colors.HexColor('#e8eef1')),('FONTNAME',(0,0),(-1,0),'Helvetica-Bold'),('FONTSIZE',(0,0),(-1,-1),7.5),('VALIGN',(0,0),(-1,-1),'TOP'),('LEFTPADDING',(0,0),(-1,-1),4),('RIGHTPADDING',(0,0),(-1,-1),4),('TOPPADDING',(0,0),(-1,-1),3),('BOTTOMPADDING',(0,0),(-1,-1),3)])); story += [Spacer(1,5),t,Spacer(1,5)]
    pdf.multiBuild(story,canvasmaker=_canvas(customer,estimate,_logo(docx_bytes))); return out.getvalue()


def render_pdf_v3(db, sow, rev):
    template=db.get(SOWTemplateVersion,sow.template_version_id)
    if not template or template.version_no<3: return _PREVIOUS_RENDER_PDF(db,sow,rev)
    docx=sow_service.verify_approved_content(db,sow,rev) if sow.status=='APPROVED' else sow_service.render_docx(db,sow,rev)
    return _review_pdf(docx,rev.customer or '',rev.estimate.estimate_number,sow.sow_date.strftime('%B %d, %Y'))

sow_service.render_pdf=render_pdf_v3
