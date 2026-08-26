from __future__ import annotations

import hashlib
import html
import io
import zipfile
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen.canvas import Canvas
from reportlab.platypus import BaseDocTemplate, Frame, PageBreak, PageTemplate, Paragraph as RP, Spacer, Table, TableStyle
from reportlab.platypus.tableofcontents import TableOfContents
from sqlalchemy.orm import Session

from . import sow_service
from .models import EstimateRevision
from .sow_models import SOW
from .sp_core_a import SMALL_PROJECT_TEMPLATE_KEYS, _config, _product_for_revision, _template_for_sow, small_project_support_hours
from .sp_core_b import appendix_included
from .sp_render_a import (
    _fill_appendix, _fill_commercial_table, _fill_hypercare_table, _filter_methodology,
    _replace_agreement, _replace_commercial_wording, _replace_deliverables, _replace_everywhere,
    _replace_key_user_count, _replace_project_objective,
)


def render_small_project_docx(db: Session, sow: SOW, rev: EstimateRevision) -> bytes:
    template = _template_for_sow(db, sow)
    if template.template_key not in SMALL_PROJECT_TEMPLATE_KEYS:
        raise ValueError("This is not a Small Project SOW template.")
    cfg = _config(db, sow)
    product = _product_for_revision(db, rev)

    doc = Document(io.BytesIO(template.content))
    _replace_everywhere(
        doc,
        {
            "<CustomerName>": rev.customer or "",
            "<99999999>": rev.estimate.estimate_number,
            "<Today>": sow.sow_date.strftime("%B %d, %Y"),
            "(Other DSI Entity)": rev.entity or "Data Systems International, Inc. dba Cloud Inventory®",
        },
    )
    _replace_agreement(doc, sow)
    _replace_project_objective(doc, sow.project_objective.strip())
    _replace_deliverables(doc, cfg)
    _filter_methodology(doc, db, sow, rev, cfg)
    _replace_key_user_count(doc, cfg.key_user_training_count)
    support_hours = small_project_support_hours(db, rev, product)
    _fill_hypercare_table(doc, sow, support_hours)
    _fill_commercial_table(doc, rev)
    _fill_appendix(doc, db, sow, rev, cfg, product)
    _replace_commercial_wording(doc, rev, sow)

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")
    for field in doc._element.findall(".//" + qn("w:fldChar")):
        if field.get(qn("w:fldCharType")) == "begin":
            field.set(qn("w:dirty"), "true")

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def small_project_content_hash_for(
    db: Session, sow: SOW, rev: EstimateRevision
) -> tuple[str, str, bytes]:
    docx = render_small_project_docx(db, sow, rev)
    text = sow_service.canonical_text(docx)
    return hashlib.sha256(text.encode("utf-8")).hexdigest(), text, docx


def verify_small_project_approved_content(
    db: Session, sow: SOW, rev: EstimateRevision
) -> bytes:
    digest, _, docx = small_project_content_hash_for(db, sow, rev)
    if sow.status == "APPROVED" and sow.content_hash and digest != sow.content_hash:
        raise ValueError(
            "The regenerated SOW no longer matches the approved wording. "
            "Historical download has been blocked for audit safety."
        )
    return docx


def _logo(docx_bytes: bytes):
    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as archive:
            names = [
                name for name in archive.namelist()
                if name.startswith("word/media/")
                and name.lower().endswith((".png", ".jpg", ".jpeg"))
            ]
            return archive.read(names[0]) if names else None
    except Exception:
        return None


class _PDFDoc(BaseDocTemplate):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._seq = 0

    def beforeDocument(self):
        self._seq = 0
        return super().beforeDocument()

    def afterFlowable(self, flowable):
        if isinstance(flowable, RP) and hasattr(flowable, "_toc_level"):
            self._seq += 1
            key = f"sow_h_{self._seq}"
            self.canv.bookmarkPage(key)
            self.notify("TOCEntry", (flowable._toc_level, flowable.getPlainText(), self.page, key))


def _pdf_canvas(customer: str, estimate: str, logo_bytes: bytes | None):
    from .sow_pdf_v3 import COVER_FOOTER

    class _Canvas(Canvas):
        def __init__(self, *args, **kwargs):
            Canvas.__init__(self, *args, **kwargs)
            self._states = []
            self._logo = ImageReader(io.BytesIO(logo_bytes)) if logo_bytes else None

        def showPage(self):
            self._states.append(dict(self.__dict__))
            self._startPage()

        def save(self):
            total = len(self._states)
            for state in self._states:
                self.__dict__.update(state)
                self._decorate(total)
                Canvas.showPage(self)
            Canvas.save(self)

        def _decorate(self, total):
            width, height = letter
            left = .72 * inch
            right = width - .72 * inch
            if self._logo:
                iw, ih = self._logo.getSize()
                w = 120
                h = w * ih / iw
                self.drawImage(
                    self._logo, left, height - .28 * inch - h,
                    width=w, height=h, mask="auto", preserveAspectRatio=True,
                )
            if self._pageNumber > 1:
                self.setFont("Helvetica-Bold", 9)
                self.drawRightString(right, height - .82 * inch, customer)
            self.setLineWidth(1)
            self.line(left, height - 1.03 * inch, right, height - 1.03 * inch)
            if self._pageNumber == 1:
                style = ParagraphStyle(
                    "coverFooter", fontName="Helvetica", fontSize=6.3, leading=7.2
                )
                paragraph = RP(html.escape(COVER_FOOTER), style)
                paragraph.wrap(right - left, .6 * inch)
                paragraph.drawOn(self, left, .28 * inch)
            else:
                self.setFont("Helvetica-Bold", 7.5)
                self.drawString(left, .42 * inch, f"Estimate Number: {estimate}")
                self.drawCentredString(
                    width / 2, .42 * inch, f"Page {self._pageNumber} of {total}"
                )

    return _Canvas


def _review_small_project_pdf(
    docx_bytes: bytes, customer: str, estimate: str, sow_date: str
) -> bytes:
    doc = Document(io.BytesIO(docx_bytes))
    out = io.BytesIO()
    frame = Frame(
        .72 * inch,
        .72 * inch,
        letter[0] - 1.44 * inch,
        letter[1] - 1.78 * inch,
        id="body",
        topPadding=0,
        bottomPadding=0,
    )
    pdf = _PDFDoc(out, pagesize=letter)
    pdf.addPageTemplates([PageTemplate(id="sow", frames=[frame])])
    styles = getSampleStyleSheet()

    def add(name, parent, **kwargs):
        styles.add(ParagraphStyle(name, parent=styles[parent], **kwargs))

    add("CoverTitle", "Title", fontName="Helvetica-Bold", fontSize=20, leading=24, alignment=TA_CENTER, spaceAfter=24)
    add("CoverCustomer", "Normal", fontName="Helvetica-Bold", fontSize=16, leading=20, alignment=TA_CENTER, spaceAfter=60)
    add("CoverSubtitle", "Normal", fontName="Helvetica-Bold", fontSize=12.5, leading=16, alignment=TA_CENTER, spaceAfter=22)
    add("CoverMeta", "Normal", fontName="Helvetica-Bold", fontSize=11, leading=14, alignment=TA_CENTER, spaceAfter=28)
    add("TOCTitle", "Heading1", fontName="Helvetica-Bold", fontSize=16, leading=20, alignment=TA_CENTER, spaceAfter=18)
    add("H1S", "Heading1", fontName="Helvetica-Bold", fontSize=14, leading=17, spaceAfter=10, keepWithNext=True)
    add("H2S", "Heading2", fontName="Helvetica-Bold", fontSize=11.5, leading=14, spaceBefore=7, spaceAfter=5, keepWithNext=True)
    add("H3S", "Heading3", fontName="Helvetica-Bold", fontSize=10.5, leading=13, spaceBefore=5, spaceAfter=4, keepWithNext=True)
    add("BodyS", "BodyText", fontName="Helvetica", fontSize=8.7, leading=11.3, spaceAfter=4)
    add("BulletS", "BodyS", leftIndent=18, firstLineIndent=-8, bulletIndent=6, spaceAfter=2)
    add("AppendixS", "Heading1", fontName="Helvetica-Bold", fontSize=16, leading=20, alignment=TA_CENTER, spaceAfter=10, keepWithNext=True)

    story = [
        Spacer(1, 1.28 * inch),
        RP("Statement Of Work", styles["CoverTitle"]),
        RP(html.escape(customer), styles["CoverCustomer"]),
        RP("Delivery Estimate for Small Projects", styles["CoverSubtitle"]),
        RP(f"Estimate Number: {html.escape(estimate)}", styles["CoverMeta"]),
        RP(f"Date: {html.escape(sow_date)}", styles["CoverMeta"]),
        PageBreak(),
        RP("Table of Contents", styles["TOCTitle"]),
    ]
    toc = TableOfContents()
    toc.levelStyles = [
        ParagraphStyle("T1", fontName="Helvetica-Bold", fontSize=9.2, leading=12, leftIndent=0, spaceBefore=4),
        ParagraphStyle("T2", fontName="Helvetica", fontSize=8.5, leading=11, leftIndent=14),
        ParagraphStyle("T3", fontName="Helvetica", fontSize=8.2, leading=10, leftIndent=28),
    ]
    story += [toc, PageBreak()]

    children = list(doc._element.body)
    toc_index = None
    for index, child in enumerate(children):
        instruction = " ".join(
            (node.text or "") for node in child.findall(".//" + qn("w:instrText"))
        )
        if "TOC" in instruction:
            toc_index = index
            break

    counters = [0, 0, 0]
    for child in children[(toc_index + 1 if toc_index is not None else 0):]:
        if child.tag == qn("w:p"):
            p = Paragraph(child, doc)
            text = p.text.strip()
            style = p.style.name if p.style else ""
            if not text:
                continue
            if text == "Appendix A":
                story.append(PageBreak())
                flow = RP("Appendix A", styles["AppendixS"])
                flow._toc_level = 0
                story.append(flow)
                continue
            if style == "Heading 1":
                counters = [counters[0] + 1, 0, 0]
                story.append(PageBreak())
                flow = RP(f"{counters[0]}.0 {html.escape(text)}", styles["H1S"])
                flow._toc_level = 0
                story.append(flow)
            elif style == "Heading 2":
                counters[1] += 1
                counters[2] = 0
                flow = RP(f"{counters[0]}.{counters[1]} {html.escape(text)}", styles["H2S"])
                flow._toc_level = 1
                story.append(flow)
            elif style == "Heading 3":
                counters[2] += 1
                flow = RP(
                    f"{counters[0]}.{counters[1]}.{counters[2]} {html.escape(text)}",
                    styles["H3S"],
                )
                flow._toc_level = 2
                story.append(flow)
            else:
                story.append(
                    RP(
                        html.escape(text),
                        styles["BulletS"] if style == "List Paragraph" else styles["BodyS"],
                        bulletText="•" if style == "List Paragraph" else None,
                    )
                )
        elif child.tag == qn("w:tbl"):
            source = DocxTable(child, doc)
            data = [
                [RP(html.escape(cell.text.strip()), styles["BodyS"]) for cell in row.cells]
                for row in source.rows
            ]
            if data:
                table = Table(data, repeatRows=1, hAlign="LEFT")
                table.setStyle(
                    TableStyle(
                        [
                            ("GRID", (0, 0), (-1, -1), .45, colors.HexColor("#7d8d94")),
                            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8eef1")),
                            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                            ("FONTSIZE", (0, 0), (-1, -1), 7.5),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ("LEFTPADDING", (0, 0), (-1, -1), 4),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                            ("TOPPADDING", (0, 0), (-1, -1), 3),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                        ]
                    )
                )
                story += [Spacer(1, 5), table, Spacer(1, 5)]

    pdf.multiBuild(
        story,
        canvasmaker=_pdf_canvas(customer, estimate, _logo(docx_bytes)),
    )
    return out.getvalue()


def render_small_project_pdf(db: Session, sow: SOW, rev: EstimateRevision) -> bytes:
    docx = (
        verify_small_project_approved_content(db, sow, rev)
        if sow.status == "APPROVED"
        else render_small_project_docx(db, sow, rev)
    )
    return _review_small_project_pdf(
        docx,
        rev.customer or "",
        rev.estimate.estimate_number,
        sow.sow_date.strftime("%B %d, %Y"),
    )
