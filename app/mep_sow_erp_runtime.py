from __future__ import annotations

import io

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

from .cip_models import PRODUCT_MEP
from .models import EstimateRevision
from .sow_models import SOW


ERP_WORDING_VERSION = 2


def _insert_after(paragraph: Paragraph, text: str, style: str = "List Paragraph") -> None:
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    created = Paragraph(element, paragraph._parent)
    try:
        created.style = style
    except KeyError:
        created.style = "Normal"
    created.add_run(text)


def _set_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _apply_connection_wording(
    content: bytes,
    sow: SOW,
    rev: EstimateRevision,
    *,
    install_mode: str | None = None,
) -> bytes:
    if int(getattr(sow, "composition_version", 1) or 1) < ERP_WORDING_VERSION:
        return content
    version = (sow.erp_version or "").strip()
    erp = (rev.erp or "").strip()
    if not version or not erp:
        return content

    doc = Document(io.BytesIO(content))
    paragraphs = list(doc.paragraphs)

    # Small Project generates a neutral MEP Installation heading from persisted scope.
    # Resolve it to the actual architecture before inserting the connection statement.
    neutral = next((p for p in paragraphs if p.text.strip() == "MEP Installation"), None)
    if neutral is not None and install_mode in {"Cloud", "On_Prem"}:
        _set_text(
            neutral,
            "MEP Cloud Installation" if install_mode == "Cloud" else "MEP On Premises Installation",
        )
        heading = neutral
    else:
        heading = next(
            (
                p for p in paragraphs
                if p.text.strip() in {"MEP Cloud Installation", "MEP On Premises Installation"}
            ),
            None,
        )

    if heading is None:
        # No MEP platform installation section survived the controlled scope filter.
        return content

    sentence = f"MEP will be connected to {erp}, version {version}."
    if any(sentence == p.text.strip() for p in doc.paragraphs):
        return content
    _insert_after(heading, sentence)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def install_mep_sow_erp_wording(core) -> None:
    """Version-gated ERP/version composition for MEP Net New and Small Project SOWs."""
    from . import sow_service
    from . import small_project_workflow
    from . import small_project_word_runtime
    from . import sp_render_b
    from .sp_core_a import _config, _product_for_revision
    from .sp_core_b import validate_small_project_finalize as original_sp_validate

    original_mep_render = sow_service.render_docx
    original_sp_render = sp_render_b.render_small_project_docx

    def render_mep_docx(db, sow: SOW, rev: EstimateRevision) -> bytes:
        content = original_mep_render(db, sow, rev)
        return _apply_connection_wording(content, sow, rev)

    def render_sp_docx(db, sow: SOW, rev: EstimateRevision) -> bytes:
        content = original_sp_render(db, sow, rev)
        if _product_for_revision(db, rev) != PRODUCT_MEP:
            return content
        cfg = _config(db, sow)
        return _apply_connection_wording(
            content,
            sow,
            rev,
            install_mode=cfg.install_mode,
        )

    def validate_sp_with_erp(db, sow: SOW, rev: EstimateRevision) -> list[str]:
        errors = list(original_sp_validate(db, sow, rev))
        if _product_for_revision(db, rev) == PRODUCT_MEP and not (sow.erp_version or "").strip():
            errors.append(
                "ERP / System Version is required for an MEP Small Project SOW. "
                "It is used to identify the system version MEP will connect to."
            )
        return errors

    # MEP Net New hash/PDF functions resolve render_docx from sow_service globals at
    # call time, so replacing this binding automatically preserves one canonical path.
    sow_service.render_docx = render_mep_docx

    # Small Project hash/PDF functions resolve the sp_render_b global at call time.
    # Patch exported/imported render bindings as well for the controlled Word path.
    sp_render_b.render_small_project_docx = render_sp_docx
    small_project_workflow.render_small_project_docx = render_sp_docx
    small_project_word_runtime.render_small_project_docx = render_sp_docx

    small_project_workflow.validate_small_project_finalize = validate_sp_with_erp
    from . import sp_routes_b
    sp_routes_b.validate_small_project_finalize = validate_sp_with_erp
