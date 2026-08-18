from __future__ import annotations

import io
from datetime import datetime

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from sqlalchemy.orm import Session

from . import sow_service
from .cip_models import PRODUCT_MEP
from .models import User
from .services.audit import record
from .sow_models import SOWTemplateVersion, SOW_TEMPLATE_MEP_NET_NEW


V1_FILENAME = "MEP_Template_NewClient_2026_14_Controlled_v1.docx"
V2_FILENAME = "MEP_Template_NewClient_2026_14_Controlled_v2.docx"
V1_REASON = "Initial controlled MEP New Client SOW template based on MEP_Template_NewClient_2026_14."
V2_REASON = (
    "Document-layout update: template-aligned cover, dynamic Heading 1/2/3 TOC, controlled page breaks, "
    "Appendix A page start, Terms and Conditions page start, and revised headers/footers with page fields."
)
COVER_FOOTER = (
    "This statement of work estimate is the property and proprietary to Data Systems International, Inc. "
    "dba Cloud inventory® and contains trade secret and confidential information and is solely for Customer’s "
    "internal use.  Without the express written consent of Cloud Inventory ®, this estimate shall not be used, "
    "reproduced, copied, disclosed, transmitted in whole or in part.  Copyright © 2026 Data Systems International, "
    "Inc. dba Cloud Inventory®.  All rights reserved."
)
_ORIGINAL_SEED = sow_service.seed_initial_sow_template


def _replace_token_across_text_nodes(paragraph: Paragraph, old: str, new: str) -> bool:
    """Replace visible token text while preserving PAGE/NUMPAGES/TOC field XML."""
    nodes = paragraph._p.findall(".//" + qn("w:t"))
    if not nodes:
        return False
    changed = False
    while True:
        values = [node.text or "" for node in nodes]
        joined = "".join(values)
        start = joined.find(old)
        if start < 0:
            break
        end = start + len(old)
        offsets = []
        pos = 0
        for value in values:
            offsets.append((pos, pos + len(value)))
            pos += len(value)
        start_idx = next(i for i, (_, hi) in enumerate(offsets) if start < hi)
        end_idx = next(i for i, (lo, hi) in enumerate(offsets) if end <= hi and end > lo)
        start_local = start - offsets[start_idx][0]
        end_local = end - offsets[end_idx][0]
        if start_idx == end_idx:
            nodes[start_idx].text = values[start_idx][:start_local] + new + values[start_idx][end_local:]
        else:
            nodes[start_idx].text = values[start_idx][:start_local] + new
            for i in range(start_idx + 1, end_idx):
                nodes[i].text = ""
            nodes[end_idx].text = values[end_idx][end_local:]
        changed = True
    return changed


def _replace_para_text_preserving_fields(paragraph: Paragraph, replacements: dict[str, str]) -> None:
    for old, new in replacements.items():
        _replace_token_across_text_nodes(paragraph, old, new)


def _page_break_paragraph():
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    return p


def _has_page_break(element) -> bool:
    return any(br.get(qn("w:type")) == "page" for br in element.findall(".//" + qn("w:br")))


def _ensure_breaks_around_toc(doc: Document) -> None:
    """The supplied template places the TOC in an SDT. Keep it isolated by page breaks."""
    body = doc._element.body
    toc = None
    for child in list(body):
        instr = " ".join((node.text or "") for node in child.findall(".//" + qn("w:instrText")))
        if "TOC" in instr:
            toc = child
            break
    if toc is None:
        raise RuntimeError("Controlled SOW template is missing its Word Table of Contents field.")
    children = list(body)
    idx = children.index(toc)
    if idx == 0 or not _has_page_break(children[idx - 1]):
        toc.addprevious(_page_break_paragraph())
    children = list(body)
    idx = children.index(toc)
    if idx + 1 >= len(children) or not _has_page_break(children[idx + 1]):
        toc.addnext(_page_break_paragraph())


def _remove_review_formatting(root) -> None:
    for hl in list(root.findall(".//" + qn("w:highlight"))):
        parent = hl.getparent()
        if parent is not None:
            parent.remove(hl)
    for shd in list(root.findall(".//" + qn("w:shd"))):
        if (shd.get(qn("w:fill")) or "").upper() == "FFFF00":
            parent = shd.getparent()
            if parent is not None:
                parent.remove(shd)


def _remove_draft_watermark(header) -> None:
    for p in list(header._element.findall(".//" + qn("w:p"))):
        text = "".join(node.text or "" for node in p.findall(".//" + qn("w:t")))
        attrs = " ".join(str(value) for el in p.iter() for value in el.attrib.values())
        if "DRAFT" in (text + " " + attrs).upper():
            parent = p.getparent()
            if parent is not None:
                parent.remove(p)


def _keep_hypercare_together(doc: Document) -> None:
    label = next((p for p in doc.paragraphs if p.text.strip() == "Supported Deployment Locations"), None)
    if label:
        label.paragraph_format.keep_with_next = True
    table = next((t for t in doc.tables if t.rows and t.rows[0].cells[0].text.strip() == "Location Description"), None)
    if not table:
        return
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            p.paragraph_format.keep_with_next = True
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))


def _build_v2_content(v1_content: bytes) -> bytes:
    doc = Document(io.BytesIO(v1_content))

    if not doc.sections:
        raise RuntimeError("Controlled SOW template has no document section.")
    doc.sections[0].different_first_page_header_footer = True
    _ensure_breaks_around_toc(doc)

    heading1 = doc.styles["Heading 1"]
    heading1.paragraph_format.page_break_before = True
    heading1.paragraph_format.keep_with_next = True
    for style_name in ("Heading 2", "Heading 3"):
        doc.styles[style_name].paragraph_format.keep_with_next = True

    appendix = next((p for p in doc.paragraphs if p.text.strip() == "Appendix A"), None)
    if not appendix:
        raise RuntimeError("Controlled SOW template is missing Appendix A.")
    appendix.paragraph_format.page_break_before = True
    appendix.paragraph_format.keep_with_next = True
    if len(doc.sections) >= 2:
        doc.sections[1].start_type = WD_SECTION_START.NEW_PAGE

    _keep_hypercare_together(doc)

    first_footer = doc.sections[0].first_page_footer
    if first_footer.paragraphs:
        first_footer.paragraphs[0].text = COVER_FOOTER
        first_footer.paragraphs[0].style = doc.styles["Footer"]
    else:
        first_footer.add_paragraph(COVER_FOOTER, style="Footer")

    roots = [doc._element]
    for section in doc.sections:
        for part in (
            section.header, section.first_page_header, section.even_page_header,
            section.footer, section.first_page_footer, section.even_page_footer,
        ):
            roots.append(part._element)
            _remove_draft_watermark(part)
    for root in roots:
        _remove_review_formatting(root)
        for fld in root.findall(".//" + qn("w:fldChar")):
            if fld.get(qn("w:fldCharType")) == "begin":
                fld.set(qn("w:dirty"), "true")

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def seed_controlled_sow_template_v2(db: Session) -> None:
    """Create v2 only over untouched bundled v1; never override an admin-managed template."""
    _ORIGINAL_SEED(db)
    rows = (
        db.query(SOWTemplateVersion)
        .filter(SOWTemplateVersion.template_key == SOW_TEMPLATE_MEP_NET_NEW)
        .order_by(SOWTemplateVersion.version_no)
        .all()
    )
    if any(row.version_no >= 2 for row in rows):
        return
    active = next((row for row in rows if row.status == "ACTIVE"), None)
    if not active or not (
        active.version_no == 1
        and active.filename == V1_FILENAME
        and active.change_reason == V1_REASON
    ):
        return

    content = _build_v2_content(active.content)
    missing = sow_service.validate_template(content)
    if missing:
        raise RuntimeError(f"Controlled SOW template v2 is missing markers: {', '.join(missing)}")
    admin = db.query(User).filter(User.username_normalized == "admin").first()
    if not admin:
        return

    now = datetime.utcnow()
    active.status = "RETIRED"
    active.retired_at = now
    row = SOWTemplateVersion(
        template_key=SOW_TEMPLATE_MEP_NET_NEW,
        label="MEP New Client SOW",
        product_type=PRODUCT_MEP,
        customer_type="Net_New",
        version_no=2,
        status="ACTIVE",
        filename=V2_FILENAME,
        content=content,
        content_sha256=sow_service.sha256_bytes(content),
        change_reason=V2_REASON,
        created_by=admin.id,
        activated_by=admin.id,
        activated_at=now,
    )
    db.add(row)
    db.flush()
    record(db, event_type="SOW_TEMPLATE_RETIRED", user_id=admin.id,
           field_name=f"SOW_TEMPLATE:{SOW_TEMPLATE_MEP_NET_NEW}:1", old_value=active.filename,
           new_value="RETIRED", reason="Superseded by controlled document-layout template v2.")
    record(db, event_type="SOW_TEMPLATE_ACTIVATED", user_id=admin.id,
           field_name=f"SOW_TEMPLATE:{SOW_TEMPLATE_MEP_NET_NEW}:2", new_value=row.filename,
           reason=row.change_reason)
    db.commit()


sow_service._replace_para_text = _replace_para_text_preserving_fields
sow_service.seed_initial_sow_template = seed_controlled_sow_template_v2
