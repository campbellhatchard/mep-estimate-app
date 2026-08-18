from __future__ import annotations

import io
import zipfile
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION_START
from fastapi.testclient import TestClient

from app.run import app
from app.database import SessionLocal
from app.models import EstimateRevision, User
from app.sow_models import SOW, SOWTemplateVersion, SOW_TEMPLATE_MEP_NET_NEW
from app.sow_service import create_sow, render_docx


def _login(client):
    response = client.post('/login', data={'username': 'Admin', 'password': 'TestPass123!'}, follow_redirects=False)
    assert response.status_code == 303


def _approved_net_new(client):
    response = client.post('/estimates/new', data={'product_type': 'MEP'}, follow_redirects=False)
    assert response.status_code == 303
    rid = int(response.headers['location'].rsplit('/', 1)[-1])
    with SessionLocal() as db:
        rev = db.get(EstimateRevision, rid)
        rev.customer = 'Layout Test Customer'
        rev.customer_type = 'Net_New'
        rev.status = 'APPROVED'
        db.commit()
    return rid


def test_controlled_sow_v1_v2_are_retained_and_v3_is_active():
    with TestClient(app):
        with SessionLocal() as db:
            rows = db.query(SOWTemplateVersion).filter(
                SOWTemplateVersion.template_key == SOW_TEMPLATE_MEP_NET_NEW
            ).order_by(SOWTemplateVersion.version_no).all()
            assert [r.version_no for r in rows[:3]] == [1, 2, 3]
            assert rows[0].status == 'RETIRED'
            assert rows[1].status == 'RETIRED'
            assert rows[1].filename.endswith('Controlled_v2.docx')
            assert rows[2].status == 'ACTIVE'
            assert rows[2].filename.endswith('Controlled_v3.docx')


def test_v2_template_has_toc_page_rules_headers_footers_and_no_review_yellow():
    """v2 remains a historical immutable template and keeps its original layout controls."""
    with TestClient(app):
        with SessionLocal() as db:
            tmpl = db.query(SOWTemplateVersion).filter(
                SOWTemplateVersion.template_key == SOW_TEMPLATE_MEP_NET_NEW,
                SOWTemplateVersion.version_no == 2,
            ).one()
            assert tmpl.status == 'RETIRED'
            content = tmpl.content
    doc = Document(io.BytesIO(content))
    assert doc.styles['Heading 1'].paragraph_format.page_break_before is True
    assert len(doc.sections) >= 2
    assert doc.sections[1].start_type == WD_SECTION_START.NEW_PAGE
    appendix = next(p for p in doc.paragraphs if p.text.strip() == 'Appendix A')
    assert appendix.paragraph_format.page_break_before is True
    first_footer = '\n'.join(p.text for p in doc.sections[0].first_page_footer.paragraphs)
    assert 'This statement of work estimate is the property and proprietary to Data Systems International, Inc.' in first_footer

    with zipfile.ZipFile(io.BytesIO(content)) as z:
        document_xml = z.read('word/document.xml').decode('utf-8')
        settings_xml = z.read('word/settings.xml').decode('utf-8')
        all_xml = '\n'.join(
            z.read(n).decode('utf-8', errors='ignore')
            for n in z.namelist() if n.startswith('word/') and n.endswith('.xml')
        )
    assert 'TOC \\o "1-3"' in document_xml
    assert 'updateFields' in settings_xml
    assert 'FFFF00' not in all_xml
    assert '>DRAFT<' not in all_xml


def test_generated_v3_preserves_page_and_numpages_fields_while_replacing_estimate_number():
    with TestClient(app) as client:
        _login(client)
        rid = _approved_net_new(client)
        with SessionLocal() as db:
            rev = db.get(EstimateRevision, rid)
            user = db.query(User).filter_by(username_normalized='admin').one()
            sow = create_sow(db, rev, user)
            assert db.get(SOWTemplateVersion, sow.template_version_id).version_no == 3
            data = render_docx(db, sow, rev)
            estimate_number = rev.estimate.estimate_number

    with zipfile.ZipFile(io.BytesIO(data)) as z:
        footer_xml = '\n'.join(
            z.read(n).decode('utf-8', errors='ignore')
            for n in z.namelist() if n.startswith('word/footer') and n.endswith('.xml')
        )
    assert estimate_number in footer_xml
    assert '[[ESTIMATE_NUMBER]]' not in footer_xml
    assert 'PAGE' in footer_xml
    assert 'NUMPAGES' in footer_xml


def test_existing_sow_template_binding_remains_historical():
    with TestClient(app) as client:
        _login(client)
        rid = _approved_net_new(client)
        with SessionLocal() as db:
            v1 = db.query(SOWTemplateVersion).filter_by(
                template_key=SOW_TEMPLATE_MEP_NET_NEW, version_no=1
            ).one()
            rev = db.get(EstimateRevision, rid)
            user = db.query(User).filter_by(username_normalized='admin').one()
            sow = SOW(
                estimate_revision_id=rev.id,
                template_version_id=v1.id,
                sow_revision_no=1,
                status='DRAFT',
                sow_date=date.today(),
                created_by=user.id,
            )
            db.add(sow)
            db.commit()
            sid = sow.id
            v1_id = v1.id
        with SessionLocal() as db:
            assert db.get(SOW, sid).template_version_id == v1_id
