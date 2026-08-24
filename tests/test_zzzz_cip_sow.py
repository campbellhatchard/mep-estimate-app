from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient

from app.run import app
from app.cip_domain import _ensure_dynamic_scope
from app.cip_models import CIPRevisionInput, CIPScopeItem
from app.database import SessionLocal
from app.models import EstimateRevision, User
from app.services.cip_calculation import recalculate_and_store as cip_recalculate_and_store
from app.sow_models import SOW, SOWTemplateVersion
from app.cip_sow import (
    CURRENT_VERSION_CATEGORY,
    CURRENT_VERSION_KEY,
    SOW_TEMPLATE_CIP_NET_NEW,
    cip_go_live_support_hours,
)


def login(client, username="Admin", password="TestPass123!"):
    response = client.post(
        "/login",
        data={"username": username, "password": password},
        follow_redirects=False,
    )
    assert response.status_code == 303


def create_cip_approver(client):
    response = client.post(
        "/admin/users/create",
        data={
            "username": "CIP SOW Reviewer",
            "email": "cip-sow-reviewer@example.com",
            "password": "CIPReviewPass123!",
            "active": "1",
            "roles": "SOW_APPROVER",
        },
        follow_redirects=False,
    )
    assert response.status_code == 303
    with SessionLocal() as db:
        return db.query(User).filter(User.username == "CIP SOW Reviewer").one().id


def create_approved_net_new_cip(client):
    response = client.post(
        "/estimates/new", data={"product_type": "CIP"}, follow_redirects=False
    )
    assert response.status_code == 303
    rid = int(response.headers["location"].rsplit("/", 1)[-1])

    with SessionLocal() as db:
        rev = db.get(EstimateRevision, rid)
        inp = db.get(CIPRevisionInput, rid)
        rev.customer = "CIP Distribution"
        rev.customer_type = "Net_New"
        rev.entity = 'Data Systems International, Inc. dba Cloud Inventory® ("Cloud Inventory")'
        rev.billing_rate = 250
        rev.project_type = "CIP Install"
        rev.erp = "JD Edwards"

        inp.project_type = "CIP Install"
        inp.deployed_over = "JD Edwards"
        inp.gateway = True
        inp.epp_install = "On Prem"
        inp.label_sites = 1
        inp.labels_required = True
        inp.label_count = 1
        inp.rest_required = True
        inp.rest_interface_count = 1
        inp.go_live_type = "Remote All"
        inp.go_live_sites = 1

        _ensure_dynamic_scope(db, rev, inp)

        desktop = (
            db.query(CIPScopeItem)
            .filter(CIPScopeItem.revision_id == rid, CIPScopeItem.category == "DESKTOP")
            .order_by(CIPScopeItem.sort_order)
            .first()
        )
        mobile = (
            db.query(CIPScopeItem)
            .filter(CIPScopeItem.revision_id == rid, CIPScopeItem.category == "MOBILE")
            .order_by(CIPScopeItem.sort_order)
            .first()
        )
        integration = (
            db.query(CIPScopeItem)
            .filter(CIPScopeItem.revision_id == rid, CIPScopeItem.category == "INTEGRATION")
            .order_by(CIPScopeItem.sort_order)
            .first()
        )
        custom = (
            db.query(CIPScopeItem)
            .filter(CIPScopeItem.revision_id == rid, CIPScopeItem.category == "CUSTOM_DESKTOP")
            .order_by(CIPScopeItem.sort_order)
            .first()
        )
        report = (
            db.query(CIPScopeItem)
            .filter(CIPScopeItem.revision_id == rid, CIPScopeItem.category == "REPORT")
            .order_by(CIPScopeItem.sort_order)
            .first()
        )
        label = db.query(CIPScopeItem).filter(
            CIPScopeItem.revision_id == rid, CIPScopeItem.category == "LABEL"
        ).one()
        rest = db.query(CIPScopeItem).filter(
            CIPScopeItem.revision_id == rid, CIPScopeItem.category == "REST"
        ).one()

        desktop.config_type = "Baseline"
        mobile.config_type = "Mod Required"
        integration.config_type = "Baseline"
        custom.config_type = "Moderate"
        custom.description = "Customer Inventory Console"
        report.config_type = "Simple"
        report.description = "Inventory Aging Report"
        label.description = "Pallet Identification Label"
        rest.description = "Inventory Availability API"
        rest.app_count = 1

        cip_recalculate_and_store(db, rev)
        db.commit()

    assert client.post(
        f"/estimate/{rid}/status/submit", follow_redirects=False
    ).status_code == 303
    assert client.post(
        f"/estimate/{rid}/status/approve", follow_redirects=False
    ).status_code == 303
    return rid


def doc_text(doc):
    values = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    return "\n".join(values)


def test_cip_sow_current_version_setting_eligibility_and_template_seed():
    with TestClient(app) as client:
        login(client)
        rid = create_approved_net_new_cip(client)

        data = client.get(
            "/data?product=CIP&category=CIP+SOW+Setting"
        )
        assert data.status_code == 200
        assert CURRENT_VERSION_CATEGORY in data.text
        assert "Current Version" in data.text

        page = client.get(f"/estimate/{rid}/sow")
        assert page.status_code == 200
        assert "CIP New Client Statement of Work" in page.text
        assert "Prepare SOW" in page.text

        created = client.post(f"/estimate/{rid}/sow/create", follow_redirects=False)
        assert created.status_code == 303
        sid = int(created.headers["location"].rsplit("/", 1)[-1])

        with SessionLocal() as db:
            sow = db.get(SOW, sid)
            template = db.get(SOWTemplateVersion, sow.template_version_id)
            assert template.template_key == SOW_TEMPLATE_CIP_NET_NEW
            assert template.product_type == "CIP"
            assert sow.mep_product_version == "Current Version"

        admin_templates = client.get("/admin/sow-templates")
        assert admin_templates.status_code == 200
        assert "MEP New Client SOW Versions" in admin_templates.text
        assert "CIP New Client SOW Versions" in admin_templates.text


def test_cip_sow_full_workflow_scope_pdf_and_approved_word():
    with TestClient(app) as client:
        login(client)
        approver_id = create_cip_approver(client)
        rid = create_approved_net_new_cip(client)

        created = client.post(f"/estimate/{rid}/sow/create", follow_redirects=False)
        assert created.status_code == 303
        sid = int(created.headers["location"].rsplit("/", 1)[-1])

        with SessionLocal() as db:
            rev = db.get(EstimateRevision, rid)
            hours = cip_go_live_support_hours(db, rev)
            assert hours > 0

        form = {
            "agreement_type": "Software as a Service Agreement",
            "invoice_frequency": "Monthly",
            "project_objective": "Deploy and configure Cloud Inventory Platform for CIP Distribution operations.",
            "barcode_printer_count": "4",
            "erp_version": "EnterpriseOne 9.2",
            "erp_base_code_version": "9.2.6",
            "erp_tools_release": "9.2.8",
            "erp_os_version": "Windows Server 2022",
            "erp_database_version": "Oracle 19c",
            "epp_product_version": "Current Version",
            "print_methods": "CIP Print Agent, RESTful API",
            "erp_deployment_model": "Customer Managed / Private Cloud",
            "hypercare_description": "Primary Distribution Center",
            "hypercare_country": "USA",
            "hypercare_support_type": "Remote",
            "hypercare_hours": str(hours),
            "device_type": "Handheld Unit",
            "device_make_model": "Zebra MC9400",
            "device_os_version": "Android 13",
        }
        assert client.post(
            f"/sow/{sid}/save", data=form, follow_redirects=False
        ).status_code == 303

        review = client.get(f"/sow/{sid}/pdf")
        assert review.status_code == 200
        assert review.content.startswith(b"%PDF")

        assert client.post(
            f"/sow/{sid}/finalize", follow_redirects=False
        ).status_code == 303
        assert client.post(
            f"/sow/{sid}/send-approval",
            data={"approver_id": approver_id},
            follow_redirects=False,
        ).status_code == 303

        client.post("/logout", follow_redirects=False)
        login(client, "CIP SOW Reviewer", "CIPReviewPass123!")

        approvals = client.get("/approvals")
        assert approvals.status_code == 200
        assert "CIP Distribution" in approvals.text
        assert "CIP" in approvals.text

        assert client.post(
            f"/sow/{sid}/approve", follow_redirects=False
        ).status_code == 303

        with SessionLocal() as db:
            sow = db.get(SOW, sid)
            assert sow.status == "APPROVED"
            assert len(sow.content_hash or "") == 64
            assert sow.approved_text_snapshot

        docx = client.get(f"/sow/{sid}/docx")
        assert docx.status_code == 200, docx.text
        doc = Document(BytesIO(docx.content))
        text = doc_text(doc)

        assert "CIP Distribution" in text
        assert "Current Version" in text
        assert "Cloud Inventory® Managed / Public Cloud" in text
        assert "Customer Inventory Console" in text
        assert "Inventory Aging Report" in text
        assert "Pallet Identification Label" in text
        assert "Inventory Availability API" in text
        assert "JD Edwards EnterpriseOne 9.2" in text
        assert "9.2.6" in text and "9.2.8" in text
        assert "Cloud Connect Gateway" in text
        assert "Limited Load Test" not in text
        assert "[[" not in text
        assert "monthly basis" in text.lower()


def test_cip_sow_rejects_non_install_net_new_template_use():
    with TestClient(app) as client:
        login(client)
        response = client.post(
            "/estimates/new", data={"product_type": "CIP"}, follow_redirects=False
        )
        rid = int(response.headers["location"].rsplit("/", 1)[-1])
        with SessionLocal() as db:
            rev = db.get(EstimateRevision, rid)
            inp = db.get(CIPRevisionInput, rid)
            rev.customer = "Change Customer"
            rev.customer_type = "Net_New"
            rev.project_type = "CIP Change"
            inp.project_type = "CIP Change"
            cip_recalculate_and_store(db, rev)
            db.commit()
        assert client.post(
            f"/estimate/{rid}/status/submit", follow_redirects=False
        ).status_code == 303
        assert client.post(
            f"/estimate/{rid}/status/approve", follow_redirects=False
        ).status_code == 303

        page = client.get(f"/estimate/{rid}/sow")
        assert page.status_code == 200
        assert "Prepare SOW" not in page.text
        create = client.post(f"/estimate/{rid}/sow/create", follow_redirects=False)
        assert create.status_code == 409
