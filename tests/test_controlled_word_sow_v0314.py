from __future__ import annotations

import base64
import io
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree
from pypdf import PdfReader
from reportlab.pdfgen.canvas import Canvas

from app.sow_review_runtime import watermark_pdf
from app.sow_word_control import (
    COVER_FOOTER,
    NS,
    W_NS,
    apply_word_controls,
    prepare_word_presentation,
)


def _w(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"


def _tiny_png() -> bytes:
    return base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Y9Wl+0AAAAASUVORK5CYII="
    )


def _insert_toc_field(doc: Document) -> None:
    body = doc._element.body
    section = body.find(qn("w:sectPr"))
    sdt = OxmlElement("w:sdt")
    content = OxmlElement("w:sdtContent")
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    instruction = OxmlElement("w:instrText")
    instruction.text = ' TOC \\o "1-3" \\h \\z \\u '
    run.append(instruction)
    paragraph.append(run)
    content.append(paragraph)
    sdt.append(content)
    if section is None:
        body.append(sdt)
    else:
        section.addprevious(sdt)


def _source_docx() -> bytes:
    doc = Document()
    cover = doc.paragraphs[0]
    cover.text = "Statement Of Work"
    cover.add_run().add_picture(io.BytesIO(_tiny_png()))
    _insert_toc_field(doc)
    doc.add_paragraph("Introductory controlled wording")
    doc.add_heading("Project Objective", level=1)
    doc.add_paragraph("Project objective body")
    doc.add_heading("Deliverables", level=1)
    doc.add_heading("Software Deliverables", level=2)
    doc.add_paragraph("Deliverable body")
    settings = doc.settings._element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)
    for fld in doc._element.findall(".//" + qn("w:fldChar")):
        fld.set(qn("w:dirty"), "true")
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _review_pdf() -> bytes:
    out = io.BytesIO()
    canvas = Canvas(out, pagesize=(612, 792))
    canvas.drawString(72, 720, "Statement Of Work")
    canvas.showPage()
    canvas.drawString(72, 720, "Table of Contents")
    canvas.drawString(72, 690, "1.0 Project Objective")
    canvas.drawString(72, 670, "2.0 Deliverables")
    canvas.drawString(72, 650, "2.1 Software Deliverables")
    canvas.showPage()
    canvas.drawString(72, 720, "Introductory controlled wording")
    canvas.showPage()
    canvas.drawString(72, 720, "1.0 Project Objective")
    canvas.showPage()
    canvas.drawString(72, 720, "2.0 Deliverables")
    canvas.drawString(72, 690, "2.1 Software Deliverables")
    canvas.save()
    return out.getvalue()


def _settings(content: bytes):
    with zipfile.ZipFile(io.BytesIO(content), "r") as archive:
        return etree.fromstring(archive.read("word/settings.xml"))


def _document(content: bytes):
    with zipfile.ZipFile(io.BytesIO(content), "r") as archive:
        return etree.fromstring(archive.read("word/document.xml"))


def test_draft_word_has_controlled_presentation_without_initial_tracked_revisions():
    prepared = prepare_word_presentation(
        _source_docx(),
        pdf_bytes=_review_pdf(),
        customer="Example Customer",
        estimate="202608001",
        draft=True,
    )
    content = apply_word_controls(
        prepared,
        draft=True,
        password="Controlled-Test-Password-123!",
    )

    settings = _settings(content)
    document = _document(content)
    protection = settings.find("w:documentProtection", namespaces=NS)
    assert settings.find("w:trackRevisions", namespaces=NS) is not None
    assert settings.find("w:updateFields", namespaces=NS) is None
    assert protection is not None
    assert protection.get(_w("edit")) == "trackedChanges"
    assert protection.get(_w("enforcement")) == "1"
    assert protection.get(_w("algorithmName")) == "SHA-512"
    assert document.findall(".//w:ins", namespaces=NS) == []
    assert document.findall(".//w:del", namespaces=NS) == []
    assert all(field.get(_w("dirty")) is None for field in document.findall(".//w:fldChar", namespaces=NS))
    assert not any("TOC" in (node.text or "") for node in document.findall(".//w:instrText", namespaces=NS))

    visible_text = " ".join(node.text or "" for node in document.findall(".//w:t", namespaces=NS))
    assert "1. Project Objective" in visible_text
    assert "4" in visible_text
    assert "2. Deliverables" in visible_text
    assert "5" in visible_text

    with zipfile.ZipFile(io.BytesIO(content), "r") as archive:
        headers = [archive.read(name) for name in archive.namelist() if name.startswith("word/header")]
        footers = [archive.read(name) for name in archive.namelist() if name.startswith("word/footer")]
    assert sum(raw.count(b"PowerPlusWaterMarkObject") for raw in headers) >= 2
    assert all(b"#C0C0C0" in raw for raw in headers if b"PowerPlusWaterMarkObject" in raw)
    assert all(b"Calibri" in raw for raw in headers if b"PowerPlusWaterMarkObject" in raw)
    assert any(b"Example Customer" in raw for raw in headers)
    assert any(COVER_FOOTER.encode("utf-8")[:40] in raw for raw in footers)
    assert any(b"Estimate Number:" in raw and b"PAGE" in raw and b"NUMPAGES" in raw for raw in footers)


def test_approved_word_keeps_headers_footers_and_protection_without_draft_watermark():
    prepared = prepare_word_presentation(
        _source_docx(),
        pdf_bytes=_review_pdf(),
        customer="Approved Customer",
        estimate="202608002",
        draft=False,
    )
    content = apply_word_controls(
        prepared,
        draft=False,
        password="Approved-Control-Password-456!",
    )
    settings = _settings(content)
    protection = settings.find("w:documentProtection", namespaces=NS)
    assert settings.find("w:trackRevisions", namespaces=NS) is not None
    assert protection is not None
    assert protection.get(_w("edit")) == "trackedChanges"
    assert protection.get(_w("enforcement")) == "1"
    with zipfile.ZipFile(io.BytesIO(content), "r") as archive:
        headers = [archive.read(name) for name in archive.namelist() if name.startswith("word/header")]
    assert all(b"PowerPlusWaterMarkObject" not in raw for raw in headers)


def test_pdf_watermark_adds_draft_text_to_every_review_page():
    raw = _review_pdf()
    controlled = watermark_pdf(raw)
    reader = PdfReader(io.BytesIO(controlled))
    assert len(reader.pages) == 5
    for page in reader.pages:
        assert "DRAFT" in (page.extract_text() or "")


def test_release_wiring_and_sow_review_ui_are_v03141():
    run = Path("app/run.py").read_text(encoding="utf-8")
    base = Path("app/templates/base.html").read_text(encoding="utf-8")
    mep = Path("app/templates/sow.html").read_text(encoding="utf-8")
    cip = Path("app/templates/cip_sow.html").read_text(encoding="utf-8")
    render = Path("render.yaml").read_text(encoding="utf-8")
    control = Path("app/sow_word_control.py").read_text(encoding="utf-8")

    assert 'app.version = "0.3.14.1"' in run
    assert "install_sow_review_pdf_watermark()" in run
    assert "register_controlled_sow_word(app, core)" in run
    assert "Draft Word SOW" not in base
    for template in (mep, cip):
        assert "<h2>SOW Review</h2>" in template
        assert "Open PDF" in template
        assert "Draft Word SOW" in template
        assert "/sow/{{sow.id}}/docx" in template
    assert "SOW_TRACK_CHANGES_PASSWORD" in render
    assert "sync: false" in render
    assert "Controlled Word download is unavailable because" in control
    assert "prepare_word_presentation" in control
    assert "updateFields" in control
    assert "MEP_NET_NEW" in control
    assert "SOW_TEMPLATE_CIP_NET_NEW" in control
    assert "small_project" not in control.casefold()
