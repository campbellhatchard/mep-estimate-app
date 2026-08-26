"""v0.3.15 — Small Project SOW Foundation regression coverage."""

from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

from app.run import app
from app.database import SessionLocal
from app.sow_models import SOWTemplateVersion
from app.small_project_sow import (
    ASSET_DIR,
    SMALL_PROJECT_REQUIRED_PLACEHOLDERS,
    SMALL_PROJECT_TEMPLATE_KEYS,
    SOW_TEMPLATE_CIP_SMALL_PROJECT,
    SOW_TEMPLATE_MEP_SMALL_PROJECT,
    load_small_project_template_asset,
    small_project_template_meta,
    validate_small_project_template,
)


def login(client, username="Admin", password="TestPass123!"):
    response = client.post(
        "/login",
        data={"username": username, "password": password},
        follow_redirects=False,
    )
    assert response.status_code == 303


def _placeholder_docx_bytes(*, include_all=True) -> bytes:
    doc = Document()
    doc.add_paragraph("Statement of Work for <CustomerName>")
    if include_all:
        doc.add_paragraph("Estimate <99999999>")
        doc.add_paragraph("Date: <Today>")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_small_project_assets_are_real_docx_files_not_fragmented_base64():
    assert ASSET_DIR.is_dir()
    assert list(ASSET_DIR.glob("*.b64.part*")) == []
    assert list(ASSET_DIR.glob("*.b64.segment*")) == []
    for template_key in SMALL_PROJECT_TEMPLATE_KEYS:
        meta = small_project_template_meta(template_key)
        path = ASSET_DIR / meta["filename"]
        assert path.is_file()
        assert path.suffix.lower() == ".docx"


def test_bundled_small_project_assets_verify_byte_for_byte():
    import hashlib

    for template_key in SMALL_PROJECT_TEMPLATE_KEYS:
        meta = small_project_template_meta(template_key)
        content = load_small_project_template_asset(template_key)
        assert hashlib.sha256(content).hexdigest() == meta["sha256"]
        # Prove the repository asset is a valid DOCX package.
        Document(BytesIO(content))


def test_bundled_small_project_assets_contain_required_placeholders():
    for template_key in SMALL_PROJECT_TEMPLATE_KEYS:
        content = load_small_project_template_asset(template_key)
        assert validate_small_project_template(content) == []


def test_validate_small_project_template_reports_missing_placeholders():
    complete = _placeholder_docx_bytes(include_all=True)
    assert validate_small_project_template(complete) == []

    incomplete = _placeholder_docx_bytes(include_all=False)
    missing = validate_small_project_template(incomplete)
    assert missing == sorted(
        marker
        for marker in SMALL_PROJECT_REQUIRED_PLACEHOLDERS
        if marker != "<CustomerName>"
    )


def test_small_project_template_metadata_is_install_base_and_product_isolated():
    mep = small_project_template_meta(SOW_TEMPLATE_MEP_SMALL_PROJECT)
    cip = small_project_template_meta(SOW_TEMPLATE_CIP_SMALL_PROJECT)
    assert mep["customer_type"] == "Install_Base"
    assert cip["customer_type"] == "Install_Base"
    assert mep["product_type"] == "MEP"
    assert cip["product_type"] == "CIP"


def test_small_project_templates_are_seeded_active_on_startup():
    with TestClient(app):
        with SessionLocal() as db:
            for template_key in SMALL_PROJECT_TEMPLATE_KEYS:
                meta = small_project_template_meta(template_key)
                row = (
                    db.query(SOWTemplateVersion)
                    .filter(SOWTemplateVersion.template_key == template_key)
                    .one()
                )
                assert row.version_no == 1
                assert row.status == "ACTIVE"
                assert row.customer_type == "Install_Base"
                assert row.product_type == meta["product_type"]
                assert row.content_sha256 == meta["sha256"]
                assert row.filename == meta["filename"]


def test_admin_sow_templates_page_shows_all_four_families_and_preserves_net_new():
    with TestClient(app) as client:
        login(client)
        response = client.get("/admin/sow-templates")
        assert response.status_code == 200
        assert "MEP New Client SOW Versions" in response.text
        assert "CIP New Client SOW Versions" in response.text
        assert "MEP Small Project SOW Versions" in response.text
        assert "CIP Small Project SOW Versions" in response.text


def test_upload_rejects_unknown_template_key():
    with TestClient(app) as client:
        login(client)
        response = client.post(
            "/admin/sow-templates/upload",
            data={"change_reason": "regression check", "template_key": "NOT_A_REAL_KEY"},
            files={
                "file": (
                    "template.docx",
                    _placeholder_docx_bytes(),
                    "application/octet-stream",
                )
            },
        )
        assert response.status_code == 400
        assert "Unknown SOW template type." in response.text


def test_upload_small_project_draft_requires_placeholders_then_succeeds():
    with TestClient(app) as client:
        login(client)

        bad = client.post(
            "/admin/sow-templates/upload",
            data={
                "change_reason": "missing placeholder regression check",
                "template_key": SOW_TEMPLATE_MEP_SMALL_PROJECT,
            },
            files={
                "file": (
                    "MEP_Small_Project_Draft.docx",
                    _placeholder_docx_bytes(include_all=False),
                    "application/octet-stream",
                )
            },
            follow_redirects=False,
        )
        assert bad.status_code == 400

        good = client.post(
            "/admin/sow-templates/upload",
            data={
                "change_reason": "Small Project foundation regression coverage",
                "template_key": SOW_TEMPLATE_CIP_SMALL_PROJECT,
            },
            files={
                "file": (
                    "CIP_Small_Project_Draft.docx",
                    _placeholder_docx_bytes(include_all=True),
                    "application/octet-stream",
                )
            },
            follow_redirects=False,
        )
        assert good.status_code == 303

        with SessionLocal() as db:
            draft = (
                db.query(SOWTemplateVersion)
                .filter(
                    SOWTemplateVersion.template_key == SOW_TEMPLATE_CIP_SMALL_PROJECT,
                    SOWTemplateVersion.version_no == 2,
                )
                .one()
            )
            assert draft.status == "DRAFT"
            assert draft.customer_type == "Install_Base"
            assert draft.product_type == "CIP"
