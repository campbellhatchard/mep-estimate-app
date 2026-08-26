"""Small Project SOW authoring/workflow regression coverage."""

from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient

from app.run import app
from app.database import SessionLocal
from app.cip_models import CIPRevisionInput
from app.models import EstimateRevision
from app.sow_models import SOW
from app.small_project_models import SmallProjectSOWConfig
from app.small_project_sow import (
    SOW_TEMPLATE_CIP_SMALL_PROJECT,
    SOW_TEMPLATE_MEP_SMALL_PROJECT,
)
from app.small_project_workflow import (
    WEEKEND_HOLIDAY_CLAUSE,
    render_small_project_docx,
    small_project_estimate_eligible,
)


def login(client, username="Admin", password="TestPass123!"):
    response = client.post(
        "/login",
        data={"username": username, "password": password},
        follow_redirects=False,
    )
    assert response.status_code == 303


def _new_estimate(client, product):
    response = client.post(
        "/estimates/new",
        data={"product_type": product},
        follow_redirects=False,
    )
    assert response.status_code == 303
    return int(response.headers["location"].rstrip("/").split("/")[-1])


def _document_text(content: bytes) -> str:
    doc = Document(BytesIO(content))
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def test_small_project_eligibility_is_install_base_approved_and_project_specific():
    with TestClient(app) as client:
        login(client)
        rid = _new_estimate(client, "MEP")
        with SessionLocal() as db:
            rev = db.get(EstimateRevision, rid)
            rev.status = "APPROVED"
            rev.customer_type = "Install_Base"
            rev.project_type = "Small Project"
            db.commit()
            assert small_project_estimate_eligible(db, rev)

            rev.customer_type = "Net_New"
            db.commit()
            assert not small_project_estimate_eligible(db, rev)

            rev.customer_type = "Install_Base"
            rev.project_type = "MEP Cloud"
            db.commit()
            assert not small_project_estimate_eligible(db, rev)


def test_mep_small_project_authoring_uses_controlled_template_and_removes_examples(monkeypatch):
    with TestClient(app) as client:
        login(client)
        rid = _new_estimate(client, "MEP")
        with SessionLocal() as db:
            rev = db.get(EstimateRevision, rid)
            rev.status = "APPROVED"
            rev.customer_type = "Install_Base"
            rev.project_type = "Small Project"
            rev.customer = "Small Project Regression Customer"
            rev.entity = "Data Systems International, Inc. dba Cloud Inventory®"
            rev.billing_rate = 250
            rev.currency = "US Dollars"
            rev.calculated_hours = 12
            rev.calculated_fees = 3000
            rev.go_live_sites = 0
            rev.go_live_type = "None"
            db.commit()

        empty = client.get(f"/estimate/{rid}/sow")
        assert empty.status_code == 200
        assert "Ready to Author" in empty.text

        created = client.post(
            f"/estimate/{rid}/sow/create",
            follow_redirects=False,
        )
        assert created.status_code == 303
        sid = int(created.headers["location"].rstrip("/").split("/")[-1])

        with SessionLocal() as db:
            sow = db.get(SOW, sid)
            rev = db.get(EstimateRevision, rid)
            assert sow.template_version.template_key == SOW_TEMPLATE_MEP_SMALL_PROJECT
            cfg = (
                db.query(SmallProjectSOWConfig)
                .filter(SmallProjectSOWConfig.sow_id == sow.id)
                .one()
            )
            baseline = next(row for row in cfg.deliverables if row.deliverable_key == "BASELINE_APPS")
            baseline.include = True
            baseline.scope_description = "Configure Cycle Count for the approved Small Project scope."
            sow.invoice_frequency = "Monthly"
            db.commit()

            content = render_small_project_docx(db, sow, rev)
            text = _document_text(content)

        assert "Small Project Regression Customer" in text
        assert "Configure Cycle Count for the approved Small Project scope." in text
        assert WEEKEND_HOLIDAY_CLAUSE in text
        assert "on a monthly basis" in text
        assert "Objective1" not in text
        assert "Put Away Application" not in text
        assert "(Other DSI Entity)" not in text
        assert "<CustomerName>" not in text
        assert "<CUSTOMERNAME>" not in text

        page = client.get(f"/sow/{sid}")
        assert page.status_code == 200
        assert "MEP Small Project Statement of Work" in page.text
        assert "Time & Materials Commercials" in page.text

        monkeypatch.setenv("SOW_TRACK_CHANGES_PASSWORD", "RegressionOnly-StrongPassword")
        word = client.get(f"/sow/{sid}/docx")
        assert word.status_code == 200
        assert word.headers["content-type"].startswith(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


def test_cip_small_project_uses_cip_family_and_has_no_mep_install_mode():
    with TestClient(app) as client:
        login(client)
        rid = _new_estimate(client, "CIP")
        with SessionLocal() as db:
            rev = db.get(EstimateRevision, rid)
            inp = db.get(CIPRevisionInput, rid)
            assert inp is not None
            rev.status = "APPROVED"
            rev.customer_type = "Install_Base"
            rev.project_type = "Small Project"
            rev.customer = "CIP Small Project Customer"
            inp.project_type = "Small Project"
            db.commit()
            assert small_project_estimate_eligible(db, rev)

        created = client.post(
            f"/estimate/{rid}/sow/create",
            follow_redirects=False,
        )
        assert created.status_code == 303
        sid = int(created.headers["location"].rstrip("/").split("/")[-1])

        with SessionLocal() as db:
            sow = db.get(SOW, sid)
            cfg = (
                db.query(SmallProjectSOWConfig)
                .filter(SmallProjectSOWConfig.sow_id == sow.id)
                .one()
            )
            assert sow.template_version.template_key == SOW_TEMPLATE_CIP_SMALL_PROJECT
            assert cfg.install_mode == "None"

        page = client.get(f"/sow/{sid}")
        assert page.status_code == 200
        assert "CIP Small Project Statement of Work" in page.text
        assert "Existing Cloud Inventory hosted deployment" in page.text


def test_finalize_requires_selected_deliverable_and_hypercare_reconciliation():
    with TestClient(app) as client:
        login(client)
        rid = _new_estimate(client, "MEP")
        with SessionLocal() as db:
            rev = db.get(EstimateRevision, rid)
            rev.status = "APPROVED"
            rev.customer_type = "Install_Base"
            rev.project_type = "Small Project"
            rev.go_live_sites = 0
            rev.go_live_type = "None"
            db.commit()

        created = client.post(f"/estimate/{rid}/sow/create", follow_redirects=False)
        sid = int(created.headers["location"].rstrip("/").split("/")[-1])

        with SessionLocal() as db:
            sow = db.get(SOW, sid)
            cfg = (
                db.query(SmallProjectSOWConfig)
                .filter(SmallProjectSOWConfig.sow_id == sow.id)
                .one()
            )
            for row in cfg.deliverables:
                row.include = False
            db.commit()

        blocked = client.post(f"/sow/{sid}/finalize", follow_redirects=False)
        assert blocked.status_code == 400
        assert "Select at least one Small Project deliverable." in blocked.text

        with SessionLocal() as db:
            sow = db.get(SOW, sid)
            cfg = (
                db.query(SmallProjectSOWConfig)
                .filter(SmallProjectSOWConfig.sow_id == sow.id)
                .one()
            )
            baseline = next(row for row in cfg.deliverables if row.deliverable_key == "BASELINE_APPS")
            baseline.include = True
            baseline.scope_description = "Approved scoped configuration."
            # v0.3.19 makes ERP/System Version mandatory for every MEP Small Project
            # SOW so the connection target is explicit even when Appendix A is omitted.
            sow.erp_version = "9.2.8"
            db.commit()

        finalized = client.post(f"/sow/{sid}/finalize", follow_redirects=False)
        assert finalized.status_code == 303
        with SessionLocal() as db:
            assert db.get(SOW, sid).status == "FINALIZED"
