from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient

from app.run import app
from app.assumptions import EstimateAssumption
from app.database import SessionLocal
from app.models import EstimateRevision
from app.services.calculation import recalculate_and_store
from app.sow_models import SOW, SOWTemplateVersion
from app.sow_service import go_live_support_hours


def login(client, username='Admin', password='TestPass123!'):
    response = client.post('/login', data={'username': username, 'password': password}, follow_redirects=False)
    assert response.status_code == 303


def create_sow_approver(client):
    response = client.post('/admin/users/create', data={
        'username': 'SOW Reviewer',
        'email': 'sow-reviewer@example.com',
        'password': 'ReviewPass123!',
        'active': '1',
        'roles': 'SOW_APPROVER',
    }, follow_redirects=False)
    assert response.status_code == 303
    with SessionLocal() as db:
        from app.models import User
        return db.query(User).filter(User.username == 'SOW Reviewer').one().id


def create_net_new_estimate(client):
    response = client.post('/estimates/new', data={'product_type': 'MEP'}, follow_redirects=False)
    assert response.status_code == 303
    rid = int(response.headers['location'].rsplit('/', 1)[-1])
    with SessionLocal() as db:
        rev = db.get(EstimateRevision, rid)
        rev.customer = 'Acme Distribution'
        rev.customer_type = 'Net_New'
        rev.entity = 'Data Systems International, Inc. dba Cloud Inventory® ("Cloud Inventory")'
        rev.erp = 'Oracle JD Edwards E1'
        rev.go_live_type = 'Remote All'
        rev.go_live_sites = 1
        rev.billing_rate = 250
        recalculate_and_store(db, rev)
        db.commit()
    assert client.post(f'/estimate/{rid}/status/submit', follow_redirects=False).status_code == 303
    assert client.post(f'/estimate/{rid}/status/approve', follow_redirects=False).status_code == 303
    return rid


def required_sow_form(hours, invoice='Monthly'):
    return {
        'agreement_type': 'Software as a Service Agreement',
        'invoice_frequency': invoice,
        'project_objective': 'Deploy and configure Mobile Enterprise Platform for Acme operations.',
        'rest_api_required': 'on',
        'barcode_printer_count': '0',
        'erp_version': 'EnterpriseOne 9.2',
        'erp_base_code_version': '9.2.6',
        'erp_tools_release': '9.2.8',
        'erp_os_version': 'Windows Server 2022',
        'erp_database_version': 'Oracle 19c',
        'mep_product_version': 'MEP 9.5.5',
        'epp_product_version': '',
        'print_methods': '',
        'erp_deployment_model': 'Customer Managed / Private Cloud',
        'hypercare_description': 'Primary site',
        'hypercare_country': 'USA',
        'hypercare_support_type': 'Remote',
        'hypercare_hours': str(hours),
        'device_type': 'Handheld Unit',
        'device_make_model': 'Zebra MC9400',
        'device_os_version': 'Android 13',
    }


def test_sow_workflow_role_queue_rejection_revision_and_approval_lock():
    with TestClient(app) as client:
        login(client)
        assert 'SOW Approver' in client.get('/admin/users').text
        approver_id = create_sow_approver(client)
        rid = create_net_new_estimate(client)

        created = client.post(f'/estimate/{rid}/sow/create', follow_redirects=False)
        assert created.status_code == 303
        sid = int(created.headers['location'].rsplit('/', 1)[-1])

        with SessionLocal() as db:
            rev = db.get(EstimateRevision, rid)
            hours = go_live_support_hours(db, rev)
            assert hours >= 0

        saved = client.post(f'/sow/{sid}/save', data=required_sow_form(hours), follow_redirects=False)
        assert saved.status_code == 303
        assert client.post(f'/sow/{sid}/finalize', follow_redirects=False).status_code == 303
        sent = client.post(f'/sow/{sid}/send-approval', data={'approver_id': approver_id}, follow_redirects=False)
        assert sent.status_code == 303

        client.post('/logout', follow_redirects=False)
        login(client, 'SOW Reviewer', 'ReviewPass123!')
        queue = client.get('/approvals')
        assert queue.status_code == 200
        assert 'Acme Distribution' in queue.text
        assert client.get(f'/sow/{sid}/pdf').status_code == 200
        assert client.post(f'/sow/{sid}/reject', data={'reason': ''}).status_code == 400
        rejected = client.post(f'/sow/{sid}/reject', data={'reason': 'Appendix A ERP deployment detail needs clarification.'}, follow_redirects=False)
        assert rejected.status_code == 303

        client.post('/logout', follow_redirects=False)
        login(client)
        newest = client.post(f'/sow/{sid}/new-revision', follow_redirects=False)
        assert newest.status_code == 303
        sid2 = int(newest.headers['location'].rsplit('/', 1)[-1])
        assert sid2 != sid
        saved = client.post(f'/sow/{sid2}/save', data=required_sow_form(hours), follow_redirects=False)
        assert saved.status_code == 303
        assert client.post(f'/sow/{sid2}/finalize', follow_redirects=False).status_code == 303
        assert client.post(f'/sow/{sid2}/send-approval', data={'approver_id': approver_id}, follow_redirects=False).status_code == 303

        client.post('/logout', follow_redirects=False)
        login(client, 'SOW Reviewer', 'ReviewPass123!')
        assert client.post(f'/sow/{sid2}/approve', follow_redirects=False).status_code == 303
        with SessionLocal() as db:
            sow = db.get(SOW, sid2)
            assert sow.status == 'APPROVED'
            assert len(sow.content_hash or '') == 64
            assert sow.approved_text_snapshot
            assert sow.template_version_id == db.query(SOWTemplateVersion).filter(SOWTemplateVersion.status == 'ACTIVE').one().id

        docx = client.get(f'/sow/{sid2}/docx')
        assert docx.status_code == 200
        doc = Document(BytesIO(docx.content))
        text = '\n'.join(p.text for p in doc.paragraphs)
        assert 'Acme Distribution' in text
        assert '4.14 Project Specific Assumptions' not in text
        assert 'Nextworld EAP Platform Administrator' not in text
        assert 'monthly basis' in text.lower()


def test_estimate_assumptions_feed_sow_section_414():
    with TestClient(app) as client:
        login(client)
        rid = create_net_new_estimate(client)
        with SessionLocal() as db:
            db.add(EstimateAssumption(
                revision_id=rid,
                text='Customer will provide validated production-like test data before UAT.',
                sort_order=1,
            ))
            db.commit()

        created = client.post(f'/estimate/{rid}/sow/create', follow_redirects=False)
        sid = int(created.headers['location'].rsplit('/', 1)[-1])
        with SessionLocal() as db:
            hours = go_live_support_hours(db, db.get(EstimateRevision, rid))
        client.post(f'/sow/{sid}/save', data=required_sow_form(hours), follow_redirects=False)
        with SessionLocal() as db:
            from app.sow_service import render_docx
            sow = db.get(SOW, sid)
            docx_content = render_docx(db, sow, db.get(EstimateRevision, rid))
        doc = Document(BytesIO(docx_content))
        text = '\n'.join(p.text for p in doc.paragraphs)
        assert '4.14 Project Specific Assumptions' in text
        assert 'Customer will provide validated production-like test data before UAT.' in text
